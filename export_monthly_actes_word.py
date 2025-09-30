"""Export mensuel des prestations par acte (2025) vers un fichier Word.

Ce script lit le fichier source `Classeur1.xlsx`, détecte la colonne "acte" et la
colonne "date", calcule le nombre de prestations par acte pour chaque mois de
janvier à juillet 2025, ajoute une colonne Total, puis écrit le résultat dans
un document Word `Nombre_mensuel_par_acte_2025.docx`.

Le calcul suit la même logique que `analyse_prestations_full.py` (détection
des colonnes, parsing date) afin d'assurer la cohérence entre les exports.
"""
from __future__ import annotations
from pathlib import Path
import sys
import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

SOURCE_FILE = Path('Classeur1.xlsx')
OUTPUT_DOCX = Path('Nombre_mensuel_par_acte_2025.docx')

MONTH_NAMES = {
    1: 'Janvier', 2: 'Février', 3: 'Mars', 4: 'Avril', 5: 'Mai', 6: 'Juin', 7: 'Juillet'
}


def load_source(fpath: Path) -> tuple[pd.DataFrame, str | None, str | None]:
    if not fpath.exists():
        raise FileNotFoundError(f"Fichier source introuvable: {fpath}")
    df = pd.read_excel(fpath, engine='openpyxl')
    # Normaliser noms colonnes
    df.columns = [str(c).strip() for c in df.columns]
    # Détecter colonne date
    date_col = next((c for c in df.columns if str(c).lower() == 'date'), None)
    if date_col is not None:
        df[date_col] = pd.to_datetime(df[date_col], errors='coerce')
    # Détection et suppression des doublons (comportement similaire à analyse_prestations_full.py)
    try:
        dups_full = df[df.duplicated(keep=False)].copy()
        if not dups_full.empty:
            print(f"[INFO] Doublons lignes détectés: {len(dups_full)} (seront supprimés pour les analyses)")
        # Supprimer les doublons lignes avant analyses
        df.drop_duplicates(inplace=True)
    except Exception as e:
        print(f"[WARN] Impossible d'analyser/supprimer les doublons: {e}")
    # Détecter colonne acte (tolérance sur le nom)
    acte_col = next((c for c in df.columns if str(c).lower() == 'acte' or 'acte' in str(c).lower()), None)
    return df, acte_col, date_col


def compute_monthly_counts(df: pd.DataFrame, acte_col: str, date_col: str, year: int = 2025) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, dict, dict]:
    if acte_col is None:
        raise ValueError('Colonne "acte" introuvable dans le fichier source')
    if date_col is None:
        raise ValueError('Colonne "date" introuvable dans le fichier source')
    # Filtrer l'année et mois 1..7
    df = df.copy()
    df = df[pd.notna(df[date_col])]
    df['__year'] = df[date_col].dt.year
    df['__month'] = df[date_col].dt.month
    dfy = df[df['__year'] == year]
    dfy = dfy[dfy['__month'].between(1, 7)]
    # Comptage
    grouped = dfy.groupby([acte_col, '__month']).size().reset_index(name='count')
    pivot = grouped.pivot(index=acte_col, columns='__month', values='count').fillna(0).astype(int)
    # S'assurer des colonnes 1..7 dans l'ordre
    for m in range(1, 8):
        if m not in pivot.columns:
            pivot[m] = 0
    pivot = pivot.reindex(sorted(pivot.columns), axis=1)
    # Colonne Total
    pivot['Total'] = pivot.sum(axis=1)
    # Trier par Total descendant pour lisibilité
    pivot = pivot.sort_values('Total', ascending=False)
    # Remplacer en-têtes mois par noms FR
    pivot.rename(columns={m: MONTH_NAMES.get(m, str(m)) for m in range(1, 8)}, inplace=True)
    # --- Montants ---
    # Détection colonne montant et conversion numérique (logique inspirée de analyse_prestations_full.py)
    montant_col = next((c for c in df.columns if str(c).lower().startswith('montant')), None)
    if montant_col is not None:
        s = dfy[montant_col].astype(str).fillna('')
        # Nettoyage : enlever espaces insécables / espaces milliers, remplacer ',' décimal par '.'
        s = s.str.replace(r'\s+', '', regex=True).str.replace('\u00A0', '', regex=False).str.replace(',', '.', regex=False)
        montant_num = pd.to_numeric(s.replace('', pd.NA), errors='coerce').fillna(0.0)
        dfy = dfy.copy()
        dfy['_montant_num'] = montant_num.values
        grouped_amt = dfy.groupby([acte_col, '__month'])['_montant_num'].sum().reset_index(name='montant')
        pivot_amt = grouped_amt.pivot(index=acte_col, columns='__month', values='montant').fillna(0.0)
        for m in range(1, 8):
            if m not in pivot_amt.columns:
                pivot_amt[m] = 0.0
        pivot_amt = pivot_amt.reindex(sorted(pivot_amt.columns), axis=1)
        pivot_amt['Total'] = pivot_amt.sum(axis=1)
        pivot_amt = pivot_amt.reindex(index=pivot.index)  # same order as counts
        pivot_amt.rename(columns={m: MONTH_NAMES.get(m, str(m)) for m in range(1, 8)}, inplace=True)
    else:
        # aucun montant trouvé -> tableau vide montants
        pivot_amt = pd.DataFrame(0.0, index=pivot.index, columns=list(MONTH_NAMES.values()) + ['Total'])

    # --- Bénéficiaires : adhérents vs ayants droit ---
    adherent_col = next((c for c in df.columns if 'adherent_code' in str(c).lower()), None)
    beneficiaire_col = next((c for c in df.columns if str(c).lower() == 'beneficiaire'), None)
    # Détecter si un identifiant unique de bénéficiaire existe
    benef_id_col = next((c for c in df.columns if any(k in str(c).lower() for k in ('beneficiaire_id', 'beneficiairecode', 'beneficiaire_code', 'beneficiary_id', 'id_beneficiaire', 'beneficiaireid'))), None)

    # Prepare empty pivots
    cols_names = list(MONTH_NAMES.values()) + ['Total']
    pivot_benef_ad = pd.DataFrame(0, index=pivot.index, columns=cols_names)
    pivot_benef_ay = pd.DataFrame(0, index=pivot.index, columns=cols_names)

    try:
        # Normaliser texte beneficiaire si présent
        if beneficiaire_col is not None:
            typ = dfy[beneficiaire_col].astype(str).str.lower().str.replace('-', ' ').str.strip()
            is_ad = typ.str.contains('adh', na=False)
            is_ay = typ.str.contains('ayant', na=False) | typ.str.contains('ayant droit', na=False) | typ.str.contains('ayantdroit', na=False)
        else:
            # Si pas de colonne beneficiaire, on considère tout comme indifférencié
            is_ad = pd.Series([True] * len(dfy), index=dfy.index)
            is_ay = pd.Series([False] * len(dfy), index=dfy.index)

        # Choisir la méthode de comptage distinct selon les colonnes disponibles
        # Priorité: identifiant unique de bénéficiaire -> adherent_code -> composite (adherent_code + beneficiaire)
        if benef_id_col is not None:
            # On peut compter distincts pour adhérents et ayants via benef_id_col
            ad = dfy[is_ad].groupby([acte_col, '__month'])[benef_id_col].nunique().reset_index(name='ad_count')
            ay = dfy[is_ay].groupby([acte_col, '__month'])[benef_id_col].nunique().reset_index(name='ay_count')
        else:
            if adherent_col is not None:
                # Adhérents : distincts par adherent_code
                ad = dfy[is_ad].groupby([acte_col, '__month'])[adherent_col].nunique().reset_index(name='ad_count')
                # Ayants : essayer d'obtenir distincts via composite key (adherent_code + beneficiaire name)
                if beneficiaire_col is not None:
                    dfy['_ay_comp'] = dfy[adherent_col].astype(str) + '||' + dfy[beneficiaire_col].astype(str)
                    ay = dfy[is_ay].groupby([acte_col, '__month'])['_ay_comp'].nunique().reset_index(name='ay_count')
                else:
                    # pas de colonne beneficiaire -> fallback sur lignes
                    ay = dfy[is_ay].groupby([acte_col, '__month']).size().reset_index(name='ay_count')
            else:
                # pas d'adherent_code : tenter nunique sur beneficiaire_col si présent
                if beneficiaire_col is not None:
                    ad = dfy[is_ad].groupby([acte_col, '__month'])[beneficiaire_col].nunique().reset_index(name='ad_count')
                    ay = dfy[is_ay].groupby([acte_col, '__month'])[beneficiaire_col].nunique().reset_index(name='ay_count')
                else:
                    # aucun identifiant utilisable -> compter lignes
                    ad = dfy[is_ad].groupby([acte_col, '__month']).size().reset_index(name='ad_count')
                    ay = dfy[is_ay].groupby([acte_col, '__month']).size().reset_index(name='ay_count')

        # Transformer en pivots et aligner avec les actes
        if not ad.empty:
            p_ad = ad.pivot(index=acte_col, columns='__month', values='ad_count').fillna(0).astype(int)
            for m in range(1, 8):
                if m not in p_ad.columns:
                    p_ad[m] = 0
            p_ad = p_ad.reindex(sorted(p_ad.columns), axis=1)
            p_ad['Total'] = p_ad.sum(axis=1)
            p_ad.rename(columns={m: MONTH_NAMES.get(m, str(m)) for m in range(1, 8)}, inplace=True)
            p_ad = p_ad.reindex(index=pivot.index).fillna(0).astype(int)
            pivot_benef_ad.update(p_ad)

        if not ay.empty:
            p_ay = ay.pivot(index=acte_col, columns='__month', values='ay_count').fillna(0).astype(int)
            for m in range(1, 8):
                if m not in p_ay.columns:
                    p_ay[m] = 0
            p_ay = p_ay.reindex(sorted(p_ay.columns), axis=1)
            p_ay['Total'] = p_ay.sum(axis=1)
            p_ay.rename(columns={m: MONTH_NAMES.get(m, str(m)) for m in range(1, 8)}, inplace=True)
            p_ay = p_ay.reindex(index=pivot.index).fillna(0).astype(int)
            pivot_benef_ay.update(p_ay)
    except Exception:
        # en cas d'erreur, laisser pivots à zéro
        pass

    # --- Totaux 'Ensemble' (bénéficiaires distincts par mois sur l'ensemble des actes)
    overall_ad = {MONTH_NAMES[m]: 0 for m in range(1, 8)}
    overall_ay = {MONTH_NAMES[m]: 0 for m in range(1, 8)}
    try:
        # use same logic as generate_monthly_beneficiaires: group by beneficiaire and sum distinct adherent ids
        benef_id_col = next((c for c in df.columns if any(k in str(c).lower() for k in ('beneficiaire_id', 'beneficiairecode', 'beneficiaire_code', 'beneficiary_id', 'id_beneficiaire', 'beneficiaireid'))), None)
        adherent_col = next((c for c in df.columns if 'adherent_code' in str(c).lower()), None)
        beneficiaire_col = next((c for c in df.columns if str(c).lower() == 'beneficiaire'), None)
        for m in range(1, 8):
            dfm = dfy[dfy['__month'] == m]
            ad_v = 0
            ay_v = 0
            if beneficiaire_col is not None and (adherent_col is not None or benef_id_col is not None):
                # group by beneficiaire then nunique on adherent/benef id
                key = benef_id_col if benef_id_col is not None else adherent_col
                tmp = dfm.groupby(beneficiaire_col).agg({key: 'nunique'})
                idx = tmp.index.astype(str).str.lower()
                adherents_mask = idx.str.startswith(('adhérent', 'adherent'))
                ayants_mask = idx.str.contains('ayant', na=False)
                if adherents_mask.any():
                    ad_v = int(tmp.loc[adherents_mask, key].sum())
                if ayants_mask.any():
                    ay_v = int(tmp.loc[ayants_mask, key].sum())
            else:
                # fallback: distinct adherent_code across month
                if adherent_col is not None:
                    ad_v = int(dfm[adherent_col].nunique())
            overall_ad[MONTH_NAMES[m]] = ad_v
            overall_ay[MONTH_NAMES[m]] = ay_v
    except Exception:
        # leave zeros on error
        pass

    # add Totals
    overall_ad['Total'] = sum(overall_ad.values())
    overall_ay['Total'] = sum(overall_ay.values())

    return pivot, pivot_amt, pivot_benef_ad, pivot_benef_ay, overall_ad, overall_ay


def export_to_word(pivot_counts: pd.DataFrame, pivot_amounts: pd.DataFrame, pivot_benef_ad: pd.DataFrame, pivot_benef_ay: pd.DataFrame, overall_ad: dict, overall_ay: dict, out_path: Path):
    if Document is None:
        raise RuntimeError('python-docx non installé. Installez-le via `pip install python-docx`.')
    doc = Document()
    doc.add_heading('Nombre mensuel de chaque prestation (acte) - 2025', level=1)
    # Description courte
    doc.add_paragraph('Tableau : nombre de prestations par acte (ligne) et montant total par acte (ligne "Montant") pour chaque mois de janvier à juillet 2025. Ligne "Total" = somme des mois.')
    # Build table with 2 header rows: months (merged across two subcols) and subheaders Adh/Ay
    months = list(MONTH_NAMES.values())
    n_months = len(months)
    # columns: Acte | (month * 2 cols) | Total_Adh | Total_Ay
    n_cols = 1 + n_months * 2 + 2
    table = doc.add_table(rows=2, cols=n_cols)
    table.style = 'Table Grid'
    # Header row 0: Month names merged across their two subcols
    table.cell(0, 0).text = 'Acte'
    for mi, m in enumerate(months):
        c = 1 + mi * 2
        cell_top = table.cell(0, c)
        cell_top.text = m
        # merge with next cell in same row
        cell_top.merge(table.cell(0, c + 1))
    # Total header merged across last two cols
    total_c1 = 1 + n_months * 2
    table.cell(0, total_c1).text = 'Total'
    table.cell(0, total_c1).merge(table.cell(0, total_c1 + 1))
    # Header row 1: subheaders Adh / Ay
    table.cell(1, 0).text = ''
    for mi in range(n_months):
        c = 1 + mi * 2
        table.cell(1, c).text = 'Adh'
        table.cell(1, c + 1).text = 'Ay'
    table.cell(1, total_c1).text = 'Adh'
    table.cell(1, total_c1 + 1).text = 'Ay'

    # Helper to compute column index for month
    def month_col(mi):
        return 1 + mi * 2

    # Start adding data rows after header (row index 2)
    for acte in pivot_counts.index:
        # Counts row (merge month pair cells to show single count per month)
        row_counts = pivot_counts.loc[acte]
        r = table.add_row()
        r.cells[0].text = str(acte)
        for mi, m in enumerate(months):
            c = month_col(mi)
            cell = table.cell(r._index, c)
            # merge with next
            merged = cell.merge(table.cell(r._index, c + 1))
            val = int(row_counts.get(m, 0))
            merged.text = str(val)
        # merge totals into one cell
        tot_cell = table.cell(r._index, total_c1).merge(table.cell(r._index, total_c1 + 1))
        tot_cell.text = str(int(row_counts.get('Total', 0)))

        # Beneficiaries row (single row with two columns per month)
        r_b = table.add_row()
        r_b.cells[0].text = 'Bénéficiaires'
        for mi, m in enumerate(months):
            c = month_col(mi)
            adh_v = int(pivot_benef_ad.loc[acte].get(m, 0)) if acte in pivot_benef_ad.index else 0
            ay_v = int(pivot_benef_ay.loc[acte].get(m, 0)) if acte in pivot_benef_ay.index else 0
            table.cell(r_b._index, c).text = str(adh_v)
            table.cell(r_b._index, c + 1).text = str(ay_v)
        table.cell(r_b._index, total_c1).text = str(int(pivot_benef_ad.loc[acte].get('Total', 0))) if acte in pivot_benef_ad.index else '0'
        table.cell(r_b._index, total_c1 + 1).text = str(int(pivot_benef_ay.loc[acte].get('Total', 0))) if acte in pivot_benef_ay.index else '0'

        # Montant row (merge month pairs)
        r_m = table.add_row()
        r_m.cells[0].text = 'Montant'
        row_amt = pivot_amounts.loc[acte]
        for mi, m in enumerate(months):
            c = month_col(mi)
            cell = table.cell(r_m._index, c)
            merged = cell.merge(table.cell(r_m._index, c + 1))
            val = float(row_amt.get(m, 0.0))
            merged.text = f"{int(round(val)):,}".replace(',', ' ')
        tot_cell_m = table.cell(r_m._index, total_c1).merge(table.cell(r_m._index, total_c1 + 1))
        tot_cell_m.text = f"{int(round(row_amt.get('Total', 0.0))):,}".replace(',', ' ')

    # Totals rows
    # TOTAL (Prestations)
    totals_counts = [pivot_counts[m].sum() for m in months] if not pivot_counts.empty else [0] * n_months
    total_row = table.add_row()
    total_row.cells[0].text = 'TOTAL (Prestations)'
    for mi, val in enumerate(totals_counts):
        c = month_col(mi)
        merged = table.cell(total_row._index, c).merge(table.cell(total_row._index, c + 1))
        merged.text = str(int(val))
    merged_tot = table.cell(total_row._index, total_c1).merge(table.cell(total_row._index, total_c1 + 1))
    merged_tot.text = str(int(sum(totals_counts)))

    # TOTAL (Bénéficiaires) - Somme par acte: first line totals per type (somme des bénéficiaires par acte)
    totals_b_ad = [pivot_benef_ad[m].sum() for m in months]
    totals_b_ay = [pivot_benef_ay[m].sum() for m in months]
    t_b_row = table.add_row()
    t_b_row.cells[0].text = 'TOTAL (Bénéficiaires) - Somme par acte'
    for mi in range(n_months):
        c = month_col(mi)
        table.cell(t_b_row._index, c).text = str(int(totals_b_ad[mi]))
        table.cell(t_b_row._index, c + 1).text = str(int(totals_b_ay[mi]))
    table.cell(t_b_row._index, total_c1).text = str(int(sum(totals_b_ad)))
    table.cell(t_b_row._index, total_c1 + 1).text = str(int(sum(totals_b_ay)))

    # TOTAL (Bénéficiaires) - Somme par acte (Adh+Ay): combined total per month (merged cells)
    t_b_row_comb = table.add_row()
    t_b_row_comb.cells[0].text = 'TOTAL (Bénéficiaires) - Somme par acte (Adh+Ay)'
    for mi in range(n_months):
        c = month_col(mi)
        merged = table.cell(t_b_row_comb._index, c).merge(table.cell(t_b_row_comb._index, c + 1))
        val = int(totals_b_ad[mi] + totals_b_ay[mi])
        merged.text = str(val)
    merged_tot_comb = table.cell(t_b_row_comb._index, total_c1).merge(table.cell(t_b_row_comb._index, total_c1 + 1))
    merged_tot_comb.text = str(int(sum(totals_b_ad) + sum(totals_b_ay)))

    # TOTAL (Bénéficiaires) - Ensemble (distinct beneficiaries across acts)
    t_b_ens = table.add_row()
    t_b_ens.cells[0].text = 'TOTAL (Bénéficiaires) - Ensemble'
    for mi, m in enumerate(months):
        c = month_col(mi)
        val = int(overall_ad.get(m, 0) + overall_ay.get(m, 0))
        merged = table.cell(t_b_ens._index, c).merge(table.cell(t_b_ens._index, c + 1))
        merged.text = str(val)
    merged_tot_both = table.cell(t_b_ens._index, total_c1).merge(table.cell(t_b_ens._index, total_c1 + 1))
    merged_tot_both.text = str(int(overall_ad.get('Total', 0) + overall_ay.get('Total', 0)))

    # TOTAL (Montant)
    totals_amt = [pivot_amounts[m].sum() for m in months]
    t_amt_row = table.add_row()
    t_amt_row.cells[0].text = 'TOTAL (Montant)'
    for mi, val in enumerate(totals_amt):
        c = month_col(mi)
        merged = table.cell(t_amt_row._index, c).merge(table.cell(t_amt_row._index, c + 1))
        merged.text = f"{int(round(val)):,}".replace(',', ' ')
    merged_tot_amt = table.cell(t_amt_row._index, total_c1).merge(table.cell(t_amt_row._index, total_c1 + 1))
    merged_tot_amt.text = f"{int(round(sum(totals_amt))):,}".replace(',', ' ')
    # (les lignes TOTAL (Prestations) et TOTAL (Montant) ont déjà été ajoutées)
    # Basic formatting: reduce font size a bit
    style = doc.styles['Normal']
    style.font.size = Pt(10)
    out_path_parent = out_path.parent
    out_path_parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(out_path)
        print(f'Export Word: {out_path}')
    except PermissionError:
        # Fichier peut être ouvert dans Word -> essayer un nom alternatif horodaté
        from datetime import datetime
        alt = out_path.with_name(out_path.stem + '_' + datetime.now().strftime('%Y%m%d_%H%M%S') + out_path.suffix)
        try:
            doc.save(alt)
            print(f"[WARN] Impossible d'écrire '{out_path}' (peut être ouvert). Exporté vers: {alt}")
        except Exception as e:
            print(f"[ERROR] Impossible de sauvegarder le document Word: {e}")
            raise


def main():
    try:
        df, acte_col, date_col = load_source(SOURCE_FILE)
        pivot_counts, pivot_amounts, pivot_benef_ad, pivot_benef_ay, overall_ad, overall_ay = compute_monthly_counts(df, acte_col, date_col, year=2025)
        if pivot_counts.empty:
            print('Aucune prestation trouvée pour 2025 (janv.-juil.). Fichier non généré.')
            return
        export_to_word(pivot_counts, pivot_amounts, pivot_benef_ad, pivot_benef_ay, overall_ad, overall_ay, OUTPUT_DOCX)
    except Exception as e:
        print(f'[ERROR] {e}', file=sys.stderr)
        raise


if __name__ == '__main__':
    main()
