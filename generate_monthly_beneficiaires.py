"""Script pour générer le tableau des bénéficiaires par mois"""
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import calendar
import locale
from pathlib import Path

# Configuration de la locale en français
try:
    locale.setlocale(locale.LC_TIME, 'fr_FR.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'fra_fra')
    except:
        pass

# Fonction de chargement et préparation des données (inspirée de analyse_prestations_full.py)
def charger(fichier: str) -> pd.DataFrame:
    if not Path(fichier).exists():
        raise FileNotFoundError(f"Fichier introuvable: {fichier}")
    
    # Charger les données
    df = pd.read_excel(fichier, engine='openpyxl')
    
    # Normaliser noms colonnes (strip)
    df.columns = [c.strip() for c in df.columns]
    
    # Date -> datetime
    date_cols = [c for c in df.columns if c.lower() == 'date']
    if date_cols:
        dc = date_cols[0]
        df[dc] = pd.to_datetime(df[dc], errors='coerce')
        # Correction demandée : toute date postérieure à juillet 2025 -> année forcée à 2024
        mask_fix = (df[dc].dt.year == 2025) & (df[dc].dt.month > 7)
        if mask_fix.any():
            df.loc[mask_fix, dc] = df.loc[mask_fix, dc].apply(lambda d: d.replace(year=2024) if pd.notna(d) else d)
    
    # Montant -> numérique (suppression séparateurs)
    montant_cols = [c for c in df.columns if c.lower() in {'montant', 'montant demande', 'montant demandé'}]
    if montant_cols:
        mc = montant_cols[0]
        df[mc] = pd.to_numeric(df[mc].astype(str).str.replace(' ', ''), errors='coerce')
    
    # Suppression doublons lignes pour analyses
    df.drop_duplicates(inplace=True)
    
    # Filtre années 2024-2025 si date existe
    if date_cols:
        mask_annees = df[dc].dt.year.isin([2024, 2025])
        df = df[mask_annees].copy()
    
    return df

# Charger et préparer les données
df = charger('Classeur1.xlsx')

# Créer le tableau mensuel
monthly_stats = []
numero = 1  # Pour la numérotation des lignes

# Identifier les colonnes nécessaires
montant_col = next((c for c in df.columns if c.lower().startswith('montant')), None)
date_col = next((c for c in df.columns if c.lower() == 'date'), None)
adherent_col = next((c for c in df.columns if 'adherent_code' in c.lower()), None)
beneficiaire_col = next((c for c in df.columns if c.lower() == 'beneficiaire'), None)

# Préparer les données mensuelles
for year in [2024, 2025]:
    data = df[df[date_col].dt.year == year]
    for month in range(1, 13):
        month_data = data[data[date_col].dt.month == month]
        if len(month_data) > 0:
            # Calculs comme dans analyse_prestations_full.py
            montant_total = month_data[montant_col].sum()
            nb_prestations = len(month_data)
            
            # Calcul des bénéficiaires selon la logique de analyse_prestations_full.py
            if beneficiaire_col and adherent_col:
                # Répartition par type de bénéficiaire comme dans analyse_prestations_full.py
                beneficiaire_stats = month_data.groupby(beneficiaire_col).agg({
                    adherent_col: 'nunique'
                })
                
                # Identifier les types de bénéficiaires
                adherents_mask = beneficiaire_stats.index.str.lower().str.startswith(('adhérent', 'adherent'))
                ayants_droits_mask = beneficiaire_stats.index.str.lower().str.contains('ayant', case=False, na=False)
                
                # Calcul des totaux exacts
                adherents = beneficiaire_stats.loc[adherents_mask, adherent_col].sum() if adherents_mask.any() else 0
                ayants_droit = beneficiaire_stats.loc[ayants_droits_mask, adherent_col].sum() if ayants_droits_mask.any() else 0
                total_beneficiaires = adherents + ayants_droit
            
            # Formatage du mois en français
            mois_str = calendar.month_name[month].capitalize()
            
            monthly_stats.append({
                'Numéro': numero,
                'Mois': f"{mois_str} {year}",
                'Montant total': montant_total,
                'Nombre prestations': nb_prestations,
                'Nombre de bénéficiaires': total_beneficiaires,
                'Adhérents': adherents,
                'Ayants droit': ayants_droit
            })
            numero += 1

# Créer un nouveau document Word
doc = Document()
doc.add_heading('Évolution mensuelle des prestations et bénéficiaires', 0)

# Ajouter la date de génération
doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")

# Description du tableau
doc.add_paragraph('''Le détail mensuel constitue la base de vérification et de recalcul des prestations et bénéficiaires. 
Il permet d'analyser la répartition des prestations entre adhérents principaux et ayants droit, 
et de suivre l'évolution du nombre de bénéficiaires distincts au fil des mois.''')

# Créer le tableau
table = doc.add_table(rows=1, cols=7)
table.style = 'Table Grid'
table.autofit = False

# Définir les largeurs de colonnes (en pouces)
widths = [0.7, 2.0, 1.5, 1.2, 1.2, 1.2, 1.2]
for i, width in enumerate(widths):
    for cell in table.columns[i].cells:
        cell.width = Inches(width)

# Style des en-têtes
headers = ['Numéro', 'Mois', 'Montant total', 'Nombre prestations', 'Nombre de bénéficiaires', 'Adhérents', 'Ayants droit']
header_cells = table.rows[0].cells
for i, header in enumerate(headers):
    cell = header_cells[i]
    cell.text = header
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]
    run.font.bold = True
    run.font.size = Pt(10)

# Données avec formatage
for stat in monthly_stats:
    row_cells = table.add_row().cells
    
    # Numéro
    row_cells[0].text = str(stat['Numéro'])
    
    # Mois
    row_cells[1].text = stat['Mois']
    
    # Montant total (formaté avec séparateurs de milliers)
    row_cells[2].text = f"{stat['Montant total']:,.0f}".replace(",", " ")
    
    # Autres colonnes numériques
    row_cells[3].text = str(stat['Nombre prestations'])
    row_cells[4].text = str(stat['Nombre de bénéficiaires'])
    row_cells[5].text = str(stat['Adhérents'])
    row_cells[6].text = str(stat['Ayants droit'])
    
    # Appliquer la taille de police à toute la ligne
    for cell in row_cells:
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.size = Pt(10)

# Ajouter des notes explicatives
doc.add_paragraph('\nNotes et explications :')
doc.add_paragraph('''- Le nombre de bénéficiaires représente le nombre total de personnes distinctes (adhérents + ayants droit) ayant bénéficié de prestations au cours du mois.
- La colonne "Adhérents" indique le nombre d'adhérents principaux ayant bénéficié de prestations pendant le mois.
- La colonne "Ayants droit" indique le nombre de personnes à charge (conjoints, enfants, etc.) ayant bénéficié de prestations pendant le mois.
- La somme des colonnes "Adhérents" et "Ayants droit" correspond au "Nombre de bénéficiaires".
- Les montants sont exprimés en francs CFA.''')

# Sauvegarder le document avec un timestamp pour éviter les conflits
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
filename = f'rapport_prestations_beneficiaires_mensuel_{timestamp}.docx'
doc.save(filename)
print(f"Rapport généré : {filename}")