"""Analyse complète des prestations Mutuelle de Santé.
Assumptions:
- Fichier source: Classeur1.xlsx (feuille active par défaut)
- Colonnes attendues (selon échantillon):
  adherent_code, adherent_nom, adherent_prenom, adherent_genre,
  beneficiaire, identifiant_prestation, acte, date, type, sous_type,
  montant, validite, centre_nom
- Colonnes manquantes du plan (region/province, montant payé) sont absentes.
  -> On suppose region/province indisponible pour l'instant.
  -> On suppose montant_paye == montant quand validite == 'accepté'.
  -> Statut = validite.
"""
from __future__ import annotations
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from pathlib import Path
import base64
from io import BytesIO
import math
from datetime import datetime
from matplotlib.patches import Rectangle
import time

try:
    from fpdf import FPDF
    _HAS_FPDF = True
except ImportError:  # fpdf2 non installé
    _HAS_FPDF = False

try:
    from docx import Document
    from docx.shared import Inches
    _HAS_DOCX = True
except Exception as e:
    print(f"[DEBUG] Import docx échoué: {e}")
    _HAS_DOCX = False
print(f"[DEBUG] python-docx disponible: {_HAS_DOCX}")

SOURCE_FILE = 'Classeur1.xlsx'
OUTPUT_EXCEL = 'rapport_prestations_complet.xlsx'
OUTPUT_PDF = 'rapport_prestations.pdf'
OUTPUT_DOCX = 'rapport_prestations.docx'
FIG_DIR = Path('figures')
FIG_DIR.mkdir(exist_ok=True)
LOGO_PATH = Path('logo_mupol.png')  # placer votre logo ici (PNG/JPG)
# Sécurité PDF (laisser vide pour désactiver)
PDF_USER_PWD = 'Lecture2025'       # mot de passe utilisateur (ouverture)
PDF_OWNER_PWD = 'AdminSecure#2025' # mot de passe propriétaire (permissions)
PDF_PERMISSIONS = ['print']        # ex: ['print', 'copy']

# Stockage global des informations sur les doublons pour reporting
DUPLICATES_INFO: dict = {}
# Lignes dont les dates ont été corrigées (postérieures à juillet 2025 -> année 2024)
DATES_CORRIGEES_DF: pd.DataFrame | None = None
LAST_HTML_REPORT: str | None = None  # stocke le dernier HTML généré pour export Word

# Style global
sns.set_theme(style='whitegrid', context='talk', palette='Set2')
plt.rcParams.update({
    'figure.dpi': 110,
    'axes.titlesize': 14,
    'axes.labelsize': 12,
    'xtick.labelsize': 10,
    'ytick.labelsize': 10,
    'grid.color': '#d1d1d1',  # légèrement plus visible
    'grid.linewidth': 0.3,     # plus fin
    'grid.linestyle': '-',
    'axes.grid': True,
    'axes.grid.axis': 'y'
})

# ------------------ Chargement & Préparation ------------------ #

def charger(fichier: str) -> pd.DataFrame:
    if not Path(fichier).exists():
        raise FileNotFoundError(f"Fichier introuvable: {fichier}")
    df = pd.read_excel(fichier, engine='openpyxl')
    # Normaliser noms colonnes (strip)
    df.columns = [c.strip() for c in df.columns]
    # Date -> datetime
    date_cols = [c for c in df.columns if c.lower() == 'date']
    if date_cols:
        dc = date_cols[0]
        df[dc] = pd.to_datetime(df[dc], errors='coerce')
        # Correction demandée : toute date postérieure à juillet 2025 -> année forcée à 2024 (mois & jour conservés)
        try:
            global DATES_CORRIGEES_DF
            mask_fix = (df[dc].dt.year == 2025) & (df[dc].dt.month > 7)  # août (8) à décembre (12) 2025
            if mask_fix.any():
                subset = df.loc[mask_fix].copy()
                subset['date_originale'] = subset[dc]
                subset['date_corrigee'] = subset[dc].apply(lambda d: d.replace(year=2024) if pd.notna(d) else d)
                # Appliquer correction sur le dataframe principal
                df.loc[mask_fix, dc] = subset['date_corrigee']
                # Sauvegarde globale (toutes colonnes + 2 colonnes info)
                cols_order = ['date_originale','date_corrigee'] + [c for c in subset.columns if c not in {'date_originale','date_corrigee'}]
                DATES_CORRIGEES_DF = subset[cols_order]
        except Exception:
            pass
    # Montant -> numérique (suppression séparateurs)
    montant_cols = [c for c in df.columns if c.lower() in {'montant', 'montant demande', 'montant demandé'}]
    if montant_cols:
        mc = montant_cols[0]
        df[mc] = (df[mc].astype(str)
                        .str.replace('\u202f', '')
                        .str.replace('\u00a0', '')
                        .str.replace(' ', '')
                        .str.replace(',', '')
                        )
        df[mc] = pd.to_numeric(df[mc], errors='coerce')
    # Centre
    centre_cols = [c for c in df.columns if 'centre' in c.lower()]
    if centre_cols:
        cc = centre_cols[0]
        df[cc] = (df[cc].astype(str).str.strip().str.upper()
                              .str.replace(' +', ' ', regex=True))
    # Statut -> statut
    if 'validite' in [c.lower() for c in df.columns]:
        col = [c for c in df.columns if c.lower() == 'validite'][0]
        df.rename(columns={col: 'statut'}, inplace=True)
        df['statut'] = df['statut'].str.strip().str.lower()
    # Vérification et capture des doublons avant suppression
    global DUPLICATES_INFO
    # Doublons ligne entière
    dups_full = df[df.duplicated(keep=False)].copy()
    # Détection d'un identifiant unique potentiel
    id_col = next((c for c in df.columns if 'identifiant' in c.lower()), None)
    if id_col:
        dups_id = df[df.duplicated(subset=[id_col], keep=False)].copy()
    else:
        dups_id = pd.DataFrame()
    # Clé composite fréquente (si colonnes présentes)
    composite_cols = [c for c in ['adherent_code','date','acte','montant'] if c in df.columns]
    if composite_cols:
        dups_composite = df[df.duplicated(subset=composite_cols, keep=False)].copy()
    else:
        dups_composite = pd.DataFrame()
    DUPLICATES_INFO = {
        'doublons_lignes': dups_full,
        'doublons_identifiant': dups_id,
        'doublons_cle_composite': dups_composite,
        'nb_doublons_lignes': int(dups_full.shape[0]),
        'nb_doublons_identifiant': int(dups_id.shape[0]) if not dups_id.empty else 0,
        'nb_doublons_cle_composite': int(dups_composite.shape[0]) if not dups_composite.empty else 0
    }
    # Suppression doublons lignes pour analyses
    df.drop_duplicates(inplace=True)
    # Filtre années 2024-2025 si date existe
    if date_cols:
        dc = date_cols[0]
        df = df[df[dc].dt.year.isin([2024, 2025])]
    return df

# ------------------ Indicateurs ------------------ #

def compute_global(df: pd.DataFrame) -> dict:
    montant_col = next((c for c in df.columns if c.lower().startswith('montant')), None)
    date_col = next((c for c in df.columns if c.lower() == 'date'), None)
    adherent_col = next((c for c in df.columns if 'adherent_code' in c.lower()), None)
    d = {}
    if montant_col:
        d['montant_total'] = float(df[montant_col].sum())
        d['montant_moyen'] = float(df[montant_col].mean()) if len(df) else math.nan
        d['montant_mediane'] = float(df[montant_col].median()) if len(df) else math.nan
    d['nb_prestations'] = int(len(df))
    if adherent_col:
        d['nb_mutualistes_distincts'] = int(df[adherent_col].nunique())
    if 'statut' in df.columns:
        total = len(df)
        acceptes = (df['statut'] == 'accepté').sum() + (df['statut'] == 'accepte').sum()
        d['taux_acceptation_pct'] = (acceptes / total * 100) if total else math.nan
        if montant_col:
            montant_total = float(df[montant_col].sum())
            montant_paye = float(df.loc[df['statut'].isin(['accepté','accepte']), montant_col].sum())
            d['montant_paye_total'] = montant_paye
            d['montant_non_paye'] = montant_total - montant_paye
            d['pourcentage_paye_pct'] = (montant_paye / montant_total * 100) if montant_total else math.nan
    if montant_col and adherent_col and 'taux_acceptation_pct' in d:
        d['cout_moyen_mutualiste'] = float(df[montant_col].sum() / df[adherent_col].nunique()) if df[adherent_col].nunique() else math.nan
    
    # Nouvelles analyses ajoutées
    if adherent_col and montant_col:
        # Nombre de mutualistes ayant consommé
        mutualistes_ayant_consomme = int(df[adherent_col].nunique())
        d['mutualistes_ayant_consomme'] = mutualistes_ayant_consomme
        
        # Coût moyen de consommation par bénéficiaire
        # = Montant total de consommation / Nombre de mutualistes ayant consommé
        if 'montant_total' in d and mutualistes_ayant_consomme > 0:
            d['cout_moyen_par_beneficiaire'] = d['montant_total'] / mutualistes_ayant_consomme
        else:
            d['cout_moyen_par_beneficiaire'] = math.nan
        
        # Taux de recours aux prestations globales
        # (Nombre de mutualistes ayant consommé / Nombre total de mutualistes éligibles) x 100
        # Nombre total de mutualistes éligibles selon les données fournies
        nombre_total_mutualistes_eligibles = 5284
        d['taux_recours_pct'] = (mutualistes_ayant_consomme / nombre_total_mutualistes_eligibles * 100) if nombre_total_mutualistes_eligibles > 0 else math.nan
        d['nombre_total_mutualistes_eligibles'] = nombre_total_mutualistes_eligibles
    
    # Période couverte
    if date_col:
        d['periode_min'] = df[date_col].min()
        d['periode_max'] = df[date_col].max()
    return d

# ------------------ Libellés Français ------------------ #

INDIC_LABELS = {
    'montant_total': 'Montant total des prestations',
    'montant_moyen': 'Montant moyen par prestation',
    'montant_mediane': 'Montant médian par prestation',
    'nb_prestations': 'Nombre de prestations',
    'nb_mutualistes_distincts': 'Nombre de mutualistes distincts',
    'taux_acceptation_pct': "Taux d'acceptation (%)",
    'cout_moyen_mutualiste': 'Coût moyen par mutualiste',
    'mutualistes_ayant_consomme': 'Mutualistes ayant consommé',
    'cout_moyen_par_beneficiaire': 'Coût moyen de consommation par bénéficiaire',
    'taux_recours_pct': 'Taux de recours aux prestations globales (%)',
    'nombre_total_mutualistes_eligibles': 'Nombre total de mutualistes éligibles',
    'montant_paye_total': 'Montant total payé (accepté)',
    'montant_non_paye': 'Montant non payé',
    'pourcentage_paye_pct': 'Pourcentage payé (%)',
    'periode_min': 'Date première prestation',
    'periode_max': 'Date dernière prestation'
}

COL_LABELS = {
    'sum': 'Montant total',
    'count': 'Nombre de prestations',
    'mean': 'Montant moyen par prestation',
    'mutualistes_distincts': 'Nombre de mutualistes distincts'
}

HEADER_FIRST_MAP = {
    'repartition_par_acte': 'Acte',
    'repartition_par_centre': 'Centre de santé',
    'repartition_par_region': 'Région',
    'repartition_par_province': 'Province',
    'repartition_par_type': 'Type',
    'repartition_par_sous_type': 'Sous-type',
    'repartition_par_statut': 'Statut',
    'repartition_par_partenaire': 'Partenaire',
    'repartition_par_beneficiaire': 'Bénéficiaire',
    'repartition_par_genre': 'Genre'
}

# ------------------ Répartitions ------------------ #

def repartitions(df: pd.DataFrame) -> dict[str, pd.DataFrame | pd.Series]:
    out = {}
    montant_col = next((c for c in df.columns if c.lower().startswith('montant')), None)
    adherent_col = next((c for c in df.columns if 'adherent_code' in c.lower()), None)
    if not montant_col:
        return out
    def agg(col):
        base = df.groupby(col)[montant_col].agg(['sum', 'count', 'mean'])
        if adherent_col and adherent_col in df.columns:
            mut = df.groupby(col)[adherent_col].nunique().rename('mutualistes_distincts')
            base = base.join(mut)
        return base.sort_values('sum', ascending=False)
    mapping = {
        'repartition_par_acte': 'acte',
        'repartition_par_type': 'type',
        'repartition_par_sous_type': 'sous_type',
        'repartition_par_centre': next((c for c in df.columns if 'centre' in c.lower()), None),
        'repartition_par_partenaire': next((c for c in df.columns if 'parten' in c.lower()), None),
        'repartition_par_beneficiaire': next((c for c in df.columns if c.lower() == 'beneficiaire'), None),
        'repartition_par_genre': next((c for c in df.columns if 'adherent_genre' in c.lower()), None),
        'repartition_par_region': next((c for c in df.columns if c.lower() == 'region'), None),
        'repartition_par_province': None,  # géré séparément via multi-index région+province
        'repartition_par_statut': 'statut'
    }
    for key, col in mapping.items():
        if key not in ('repartition_par_province',) and col and col in df.columns:
            result = agg(col)
            # Ajout spécial pour les bénéficiaires : proportion ayants droits + ligne totale
            if key == 'repartition_par_beneficiaire' and 'mutualistes_distincts' in result.columns:
                result = result.copy()
                
                # Calculer les totaux pour adhérents et ayants droits
                adherents_mask = result.index.str.lower().str.startswith('adhérent') | result.index.str.lower().str.startswith('adherent')
                ayants_droits_mask = result.index.str.lower().str.contains('ayant', case=False, na=False)
                
                # Extraire les données pour adhérents
                if adherents_mask.any():
                    adherents_data = result.loc[adherents_mask]
                    montant_total_adherents = adherents_data['sum'].sum()
                    prestations_total_adherents = adherents_data['count'].sum()
                    mutualistes_total_adherents = adherents_data['mutualistes_distincts'].sum()
                else:
                    montant_total_adherents = prestations_total_adherents = mutualistes_total_adherents = 0
                
                # Extraire les données pour ayants droits
                if ayants_droits_mask.any():
                    ayants_droits_data = result.loc[ayants_droits_mask]
                    montant_total_ayants_droits = ayants_droits_data['sum'].sum()
                    prestations_total_ayants_droits = ayants_droits_data['count'].sum()
                    mutualistes_total_ayants_droits = ayants_droits_data['mutualistes_distincts'].sum()
                else:
                    montant_total_ayants_droits = prestations_total_ayants_droits = mutualistes_total_ayants_droits = 0
                
                # Totaux globaux
                montant_total_global = montant_total_adherents + montant_total_ayants_droits
                prestations_total_global = prestations_total_adherents + prestations_total_ayants_droits
                mutualistes_total_global = mutualistes_total_adherents + mutualistes_total_ayants_droits
                montant_moyen_global = montant_total_global / prestations_total_global if prestations_total_global > 0 else 0
                
                # Ajouter la ligne "Mutualiste" (total)
                result.loc['Mutualiste', :] = [montant_total_global, prestations_total_global, montant_moyen_global, mutualistes_total_global]
                
                # Ajouter la proportion en pourcentage
                if mutualistes_total_adherents > 0:
                    proportion_ayants_droits = (mutualistes_total_ayants_droits / mutualistes_total_adherents) * 100
                    result.loc['Proportion ayants droits/adhérents (%)', :] = [math.nan, math.nan, math.nan, proportion_ayants_droits]
            
            out[key] = result
    # Province multi-index region+province
    region_col = mapping['repartition_par_region']
    province_col = next((c for c in df.columns if c.lower() == 'province'), None)
    if region_col and province_col and region_col in df.columns and province_col in df.columns:
        grp = df.groupby([region_col, province_col])
        base = grp[montant_col].agg(['sum', 'count', 'mean'])
        if adherent_col and adherent_col in df.columns:
            mut = grp[adherent_col].nunique().rename('mutualistes_distincts')
            base = base.join(mut)
        reg_order = base.groupby(level=0)['sum'].sum().sort_values(ascending=False).index
        ordered = []
        for reg in reg_order:
            sub = base.loc[reg].sort_values('sum', ascending=False)
            for prov, row in sub.iterrows():
                ordered.append((reg, prov, *row.values))
        cols = ['Region','Province'] + list(base.columns)
        out['repartition_par_province'] = pd.DataFrame(ordered, columns=cols)
    # Top centres
    if 'repartition_par_centre' in out:
        out['top10_centres'] = out['repartition_par_centre'].head(10)
    return out

# ------------------ Analyse temporelle ------------------ #

def analyse_temporelle(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    date_col = next((c for c in df.columns if c.lower() == 'date'), None)
    montant_col = next((c for c in df.columns if c.lower().startswith('montant')), None)
    if not date_col or not montant_col:
        return pd.DataFrame(), pd.DataFrame()
    temp = df[[date_col, montant_col]].copy()
    temp['annee'] = temp[date_col].dt.year
    temp['mois'] = temp[date_col].dt.to_period('M').astype(str)
    evo = temp.groupby('mois')[montant_col].agg(['sum', 'count']).rename(columns={'sum':'montant_total','count':'nb_prestations'})
    # Comparaison T1
    temp['trimestre'] = temp[date_col].dt.to_period('Q').astype(str)
    t1 = temp[temp[date_col].dt.quarter == 1]
    comp_t1 = t1.groupby('annee')[montant_col].agg(['sum', 'count']).rename(columns={'sum':'montant_total','count':'nb_prestations'})
    return evo, comp_t1

# ------------------ Graphiques ------------------ #

def _annotate_bars(ax):
    for p in ax.patches:
        height = p.get_height()
        if not math.isnan(height):
            ax.annotate(f"{int(height):,}".replace(',', ' '),
                        (p.get_x() + p.get_width()/2, height),
                        ha='center', va='bottom', fontsize=9, rotation=0)

def plot_bar(df_series: pd.Series, title: str, filename: str, top: int | None = None, horizontal=False, pie=False, counts_series: pd.Series | None = None, annotate_vertical: bool = False):
    if df_series.empty:
        return None
    data = df_series.head(top) if top else df_series
    # Ajustement dynamique de la largeur pour lisibilité quand beaucoup de catégories verticales
    if not horizontal and not pie:
        base_w = 10
        if len(data) > 25:
            base_w = min(0.35 * len(data), 30)  # plafonner pour éviter figures énormes
        fig, ax = plt.subplots(figsize=(base_w, 6))
    else:
        fig, ax = plt.subplots(figsize=(10,5))
    if pie:
        total = data.sum()
        percentages = (data / total * 100).round(1)
        wedges, texts = ax.pie(data.values, startangle=90)
        # Pas d'autopct pour éviter chevauchement; légende externe
        legend_labels = [f"{idx} – {val:,.0f} ({pct}%)".replace(',', ' ') for idx, val, pct in zip(data.index, data.values, percentages)]
        ax.legend(wedges, legend_labels, title='Actes', loc='center left', bbox_to_anchor=(1.0, 0.5), fontsize=8)
        ax.set_ylabel('')
    else:
        if horizontal:
            sns.barplot(x=data.values, y=data.index, ax=ax)
            ax.set_xlabel('Montant total')
            ax.set_ylabel('')
            for i, (idx, v) in enumerate(zip(data.index, data.values)):
                if counts_series is not None and idx in counts_series.index:
                    c = counts_series.loc[idx]
                    ax.text(v, i, f" {int(v):,} ({int(c)})".replace(',', ' '), va='center', fontsize=9)
                else:
                    ax.text(v, i, f" {int(v):,}".replace(',', ' '), va='center', fontsize=9)
        else:
            sns.barplot(x=data.index, y=data.values, ax=ax)
            ax.set_ylabel('Montant total')
            ax.set_xlabel('')
            # Annotations sur chaque barre
            max_val = 0
            for p, idx in zip(ax.patches, data.index):
                h = p.get_height()
                if not math.isnan(h):
                    max_val = max(max_val, h)
            # Décalage et marge supérieure plus généreux pour éviter débordement
            y_offset = max_val * 0.03  # léger espace au-dessus de chaque barre
            for p, idx in zip(ax.patches, data.index):
                height = p.get_height()
                if math.isnan(height):
                    continue
                if counts_series is not None and idx in counts_series.index:
                    c = counts_series.loc[idx]
                    txt = f"{int(height):,} ({int(c)})".replace(',', ' ')
                else:
                    txt = f"{int(height):,}".replace(',', ' ')
                if annotate_vertical:
                    ax.annotate(txt,
                                (p.get_x() + p.get_width()/2, height + y_offset),
                                ha='center', va='bottom', fontsize=8, rotation=90, clip_on=False)
                else:
                    ax.annotate(txt, (p.get_x() + p.get_width()/2, height), ha='center', va='bottom', fontsize=9)
            if annotate_vertical and max_val > 0:
                ax.set_ylim(0, max_val * 1.28)  # marge verticale accrue (28%) pour contenir texte pivoté
        # Rotation dynamique si beaucoup de labels
        if not horizontal:
            rot = 90 if len(data) > 25 else 35
            ax.tick_params(axis='x', rotation=rot)
        else:
            ax.tick_params(axis='x', rotation=0)
    ax.set_title(title)
    plt.tight_layout()
    # Bordure carrée
    rect = Rectangle((0.005,0.005),0.99,0.99, transform=fig.transFigure, fill=False, lw=1.4, edgecolor='#333')
    fig.patches.append(rect)
    path = FIG_DIR / filename
    fig.savefig(path)
    plt.close(fig)
    return path

def plot_pie_group_small(series: pd.Series, title: str, filename: str, threshold: float = 0.03, donut: bool = True):
    if series.empty:
        return None
    s = series.dropna().copy()
    total = s.sum()
    if total == 0:
        return None
    props = s / total
    small = props < threshold
    if small.sum() > 1:
        s_other = s[small].sum()
        s = s[~small]
        s.loc['Autres'] = s_other
    s = s.sort_values(ascending=False)
    fig, ax = plt.subplots(figsize=(9.5,6))
    pct = (s / s.sum() * 100).round(1)
    wedgeprops = {'width':0.5} if donut else None
    wedges, _ = ax.pie(s.values, startangle=90, wedgeprops=wedgeprops, labels=None)
    legend_labels = [f"{idx} – {val:,.0f} ({p:.1f}%)".replace(',', ' ') for idx, val, p in zip(s.index, s.values, pct)]
    ax.legend(wedges, legend_labels, title='Catégories', loc='center left', bbox_to_anchor=(1.0, 0.5), fontsize=8)
    ax.set_ylabel('')
    ax.set_title(title)
    plt.tight_layout()
    # Bordure carrée
    rect = Rectangle((0.005,0.005),0.99,0.99, transform=fig.transFigure, fill=False, lw=1.4, edgecolor='#333')
    fig.patches.append(rect)
    path = FIG_DIR / filename
    fig.savefig(path)
    plt.close(fig)
    return path

def plot_region_distribution(reps: dict) -> Path | None:
    """Barres horizontales montant par région avec nombre entre parenthèses (sans courbe)."""
    key = 'repartition_par_region'
    if key not in reps:
        return None
    df = reps[key]
    if not {'sum','count'}.issubset(df.columns):
        return None
    data = df.sort_values('sum', ascending=False)
    fig, ax = plt.subplots(figsize=(10, 0.45*len(data)+1.2))
    y = range(len(data))
    ax.barh(y, data['sum'], color='#76b5c5', alpha=0.85)
    ax.set_yticks(list(y))
    ax.set_yticklabels(data.index)
    ax.invert_yaxis()  # plus grand en haut
    ax.set_xlabel('Montant total')
    ax.set_ylabel('Région')
    for yi, (v, n) in enumerate(zip(data['sum'], data['count'])):
        val_str = f"{int(v):,}".replace(',', ' ')
        ax.text(v, yi, f" {val_str} ({int(n)})", va='center', fontsize=8)
    ax.set_title('Répartition par région – Montant (Nombre)')
    fig.tight_layout()
    rect = Rectangle((0.005,0.005),0.99,0.99, transform=fig.transFigure, fill=False, lw=1.4, edgecolor='#333')
    fig.patches.append(rect)
    path = FIG_DIR / 'repartition_region_montant_nombre.png'
    fig.savefig(path)
    plt.close(fig)
    return path

def plot_province_distribution(reps: dict, top: int = 20) -> Path | None:
    """Barres horizontales provinces (top N) montant + annotations nombre."""
    key = 'repartition_par_province'
    if key not in reps:
        return None
    df = reps[key].copy()
    if 'Province' not in df.columns or 'sum' not in df.columns or 'count' not in df.columns:
        return None
    df_top = df.sort_values('sum', ascending=False).head(top)
    fig, ax = plt.subplots(figsize=(11, 0.45*len(df_top)+1.5))
    y = range(len(df_top))
    ax.barh(y, df_top['sum'], color='#a5d296', alpha=0.9)
    ax.set_yticks(list(y))
    ax.set_yticklabels(df_top['Province'])
    ax.invert_yaxis()  # plus grand en haut
    ax.set_xlabel('Montant total')
    ax.set_title(f'Repartition par province – Top {len(df_top)} (Montant (Nombre))')
    for yi, (v, n) in enumerate(zip(df_top['sum'], df_top['count'])):
        val_str = f"{int(v):,}".replace(',', ' ')
        ax.text(v, yi, f" {val_str} ({int(n)})", va='center', fontsize=8)
    fig.tight_layout()
    rect = Rectangle((0.005,0.005),0.99,0.99, transform=fig.transFigure, fill=False, lw=1.4, edgecolor='#333')
    fig.patches.append(rect)
    path = FIG_DIR / 'repartition_province_montant_nombre.png'
    fig.savefig(path)
    plt.close(fig)
    return path

def annotate_line_no_overlap(ax, x_vals, y_vals, texts, min_gap_frac=0.025, x_offset=0.35):
    """Annotate line points with labels placed to the right, clustering to avoid vertical overlap.
    - min_gap_frac: minimal vertical gap between labels as fraction of data y-range.
    - x_offset: horizontal offset (data units) to the right of each point.
    The function expands y-limits if needed and draws connector lines for shifted labels."""
    if not x_vals:
        return
    y_min, y_max = float(min(y_vals)), float(max(y_vals))
    y_range = (y_max - y_min) or 1.0
    min_gap = y_range * min_gap_frac
    # Build clusters of indices whose y are within min_gap when sorted
    sorted_idx = sorted(range(len(y_vals)), key=lambda i: y_vals[i])
    clusters = []
    current = [sorted_idx[0]]
    for idx in sorted_idx[1:]:
        if abs(y_vals[idx] - y_vals[current[-1]]) <= min_gap:
            current.append(idx)
        else:
            clusters.append(current)
            current = [idx]
    clusters.append(current)
    adjusted_y = [None]*len(y_vals)
    for cluster in clusters:
        if len(cluster) == 1:
            adjusted_y[cluster[0]] = y_vals[cluster[0]]
            continue
        # Spread cluster labels vertically centered around their mean
        orig_vals = [y_vals[i] for i in cluster]
        center = sum(orig_vals)/len(orig_vals)
        span = min_gap * (len(cluster)-1)
        start = center - span/2
        for pos, i in enumerate(cluster):
            adjusted_y[i] = start + pos*min_gap
    # Possibly extend y-limits if we pushed above
    new_y_max = max(adjusted_y + [y_max])
    if new_y_max > y_max:
        ax.set_ylim(top=new_y_max + min_gap)
    # Adjust x-limits to give space on right
    x_max = max(x_vals)
    ax.set_xlim(-0.5, x_max + 1.2)  # extra room
    # Draw annotations
    for i, (x, orig_y, lab_y, txt) in enumerate(zip(x_vals, y_vals, adjusted_y, texts)):
        ax.annotate(txt, (x + x_offset, lab_y), ha='left', va='center', fontsize=8)
        if abs(lab_y - orig_y) > 1e-9 or x_offset != 0:
            ax.plot([x, x + x_offset*0.9], [orig_y, lab_y], color='gray', linewidth=0.5, linestyle=':')

def format_abbrev(val: float) -> str:
    """Format large numbers with French style abbreviations (k, M, Md)."""
    if pd.isna(val):
        return ''
    abs_v = abs(val)
    if abs_v >= 1_000_000_000:
        return f"{val/1_000_000_000:.2f} Md".replace('.', ',')
    if abs_v >= 1_000_000:
        return f"{val/1_000_000:.2f} M".replace('.', ',')
    if abs_v >= 10_000:
        return f"{val/1_000:.1f} k".replace('.', ',')
    return f"{int(val):,}".replace(',', ' ')
def annotate_line_no_overlap(ax, x_vals, y_vals, texts, min_gap_frac=0.02, base_x_offset=0.35):
    """Advanced annotation: multi-line, adaptive staggering horizontally & vertically.
    texts: list of strings (can contain '\n')."""
    if not x_vals:
        return
    y_min, y_max = float(min(y_vals)), float(max(y_vals))
    y_range = (y_max - y_min) or 1.0
    min_gap = y_range * min_gap_frac
    # Sort points by y
    order = sorted(range(len(y_vals)), key=lambda i: y_vals[i])
    adjusted_y = [None]*len(y_vals)
    clusters = []
    cluster = [order[0]]
    for i in order[1:]:
        if abs(y_vals[i] - y_vals[cluster[-1]]) <= min_gap:
            cluster.append(i)
        else:
            clusters.append(cluster)
            cluster = [i]
    clusters.append(cluster)
    for cl in clusters:
        if len(cl) == 1:
            adjusted_y[cl[0]] = y_vals[cl[0]]
        else:
            center = sum(y_vals[k] for k in cl)/len(cl)
            span = min_gap * max(1, (len(cl)-1)) * 1.25
            start = center - span/2
            step = span / (len(cl)-1)
            for idx,pos in enumerate(cl):
                adjusted_y[pos] = start + idx*step
    # Extend ylim if necessary
    new_top = max(adjusted_y + [y_max])
    if new_top > y_max:
        ax.set_ylim(top=new_top + min_gap*0.8)
    x_max = max(x_vals)
    ax.set_xlim(-0.5, x_max + 2.2)
    # Draw annotations
    for cl in clusters:
        for idx, point_idx in enumerate(cl):
            x = x_vals[point_idx]
            oy = y_vals[point_idx]
            ly = adjusted_y[point_idx]
            # Horizontal offset with slight incremental shift to create a fan
            local_off = base_x_offset + (idx * 0.15 if len(cl) > 2 else idx * 0.1)
            ax.annotate(texts[point_idx], (x + local_off, ly), ha='left', va='center', fontsize=8, linespacing=1.0)
            if abs(ly - oy) > 1e-9 or local_off != 0:
                ax.plot([x, x + local_off*0.85], [oy, ly], color='gray', linewidth=0.5, linestyle=':')

def plot_monthly_lines(pivot_data: pd.DataFrame, title: str, filename: str, year: int) -> Path | None:
    """Crée un graphique en lignes multiples pour l'évolution mensuelle des prestations par acte."""
    try:
        if pivot_data.empty:
            return None
        
        # Préparer les données (exclure la colonne Total)
        data_cols = [col for col in pivot_data.columns if col != 'Total']
        plot_data = pivot_data[data_cols]
        
        # Garder seulement les top 10 actes les plus fréquents pour la lisibilité
        top_actes = pivot_data.nlargest(10, 'Total').index
        plot_data = plot_data.loc[top_actes]
        
        fig, ax = plt.subplots(figsize=(12, 8))
        
        # Palette de couleurs distinctes
        colors = plt.cm.tab10(np.linspace(0, 1, len(plot_data)))
        
        # Tracer une ligne par acte
        for i, (acte, row) in enumerate(plot_data.iterrows()):
            # Filtrer les valeurs non nulles pour éviter les lignes brisées
            valid_data = [(j, val) for j, val in enumerate(row) if val > 0]
            if valid_data:
                x_vals, y_vals = zip(*valid_data)
                ax.plot(x_vals, y_vals, marker='o', linewidth=2.5, markersize=6,
                       label=acte, color=colors[i], alpha=0.8)
        
        # Personnalisation du graphique
        ax.set_xlabel('Mois', fontsize=12, fontweight='bold')
        ax.set_ylabel('Nombre de prestations', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        
        # Labels des mois sur l'axe X
        ax.set_xticks(range(len(data_cols)))
        ax.set_xticklabels([col[:3].title() for col in data_cols], rotation=45)
        
        # Grille pour faciliter la lecture
        ax.grid(True, alpha=0.3, linestyle='-', linewidth=0.5)
        ax.set_facecolor('#fafafa')
        
        # Légende
        ax.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize=9)
        
        # Format des valeurs sur l'axe Y
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'{int(x):,}'.replace(',', ' ')))
        
        plt.tight_layout()
        
        # Sauvegarde
        path = FIG_DIR / filename
        fig.savefig(path, dpi=300, bbox_inches='tight')
        plt.close(fig)
        return path
        
    except Exception as e:
        print(f"Erreur lors de la création du graphique {filename}: {e}")
        return None

def plot_montant_moyen_comparison(data, title, filename):
    """Créer un graphique en barres groupées pour comparer les montants moyens par acte entre 2024 et 2025."""
    try:
        fig, ax = plt.subplots(figsize=(14, 8))
        
        # Préparer les données
        actes = [row['Acte'] for row in data]
        montants_2024 = [row['Montant_moyen_2024'] for row in data]
        montants_2025 = [row['Montant_moyen_2025'] for row in data]
        
        # Limiter à 15 actes max pour la lisibilité
        if len(actes) > 15:
            actes = actes[:15]
            montants_2024 = montants_2024[:15]
            montants_2025 = montants_2025[:15]
        
        # Configuration des barres
        x = np.arange(len(actes))
        width = 0.35
        
        # Créer les barres
        bars1 = ax.bar(x - width/2, montants_2024, width, label='2024', 
                      color='#2E86AB', alpha=0.8, edgecolor='white', linewidth=0.8)
        bars2 = ax.bar(x + width/2, montants_2025, width, label='2025', 
                      color='#A23B72', alpha=0.8, edgecolor='white', linewidth=0.8)
        
        # Ajouter les valeurs sur les barres
        def add_value_labels(bars, values):
            for bar, val in zip(bars, values):
                if val > 0:
                    height = bar.get_height()
                    ax.annotate(f'{format_abbrev(val)}',
                              xy=(bar.get_x() + bar.get_width() / 2, height),
                              xytext=(0, 3),  # 3 points de décalage vertical
                              textcoords="offset points",
                              ha='center', va='bottom', fontsize=9, fontweight='bold')
        
        add_value_labels(bars1, montants_2024)
        add_value_labels(bars2, montants_2025)
        
        # Personnalisation
        ax.set_xlabel('Types de prestations', fontsize=12, fontweight='bold')
        ax.set_ylabel('Montant moyen (FCFA)', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        
        # Rotation des labels d'actes pour la lisibilité
        ax.set_xticks(x)
        ax.set_xticklabels(actes, rotation=45, ha='right', fontsize=10)
        
        # Légende
        ax.legend(fontsize=11, loc='upper right')
        
        # Grille légère
        ax.grid(True, alpha=0.3, axis='y')
        ax.set_axisbelow(True)
        
        # Formatage de l'axe Y
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: format_abbrev(x)))
        
        plt.tight_layout()
        
        # Sauvegarde
        path = FIG_DIR / filename
        fig.savefig(path, dpi=300, bbox_inches='tight')
        plt.close(fig)
        return path
        
    except Exception as e:
        print(f"Erreur lors de la création du graphique {filename}: {e}")
        return None

def plot_proportions_prestations(proportions_data, title, filename):
    """
    Créer un graphique en barres horizontales pour les proportions de prestations (hors pharmacie)
    """
    try:
        # Prendre les 15 premiers actes pour la lisibilité
        data_top = proportions_data[:15]
        
        # Extraire les données
        actes = [row['Acte'] for row in data_top]
        proportions = [row['Proportion_pct'] for row in data_top]
        
        # Créer la figure
        fig, ax = plt.subplots(figsize=(12, 8))
        
        # Créer le graphique en barres horizontales
        colors = plt.cm.Set3(np.linspace(0, 1, len(actes)))
        bars = ax.barh(range(len(actes)), proportions, color=colors, alpha=0.8, edgecolor='black', linewidth=0.5)
        
        # Ajouter les valeurs sur les barres
        for i, (bar, value) in enumerate(zip(bars, proportions)):
            width = bar.get_width()
            ax.text(width + 0.1, bar.get_y() + bar.get_height()/2, 
                   f'{value:.1f}%', ha='left', va='center', fontweight='bold', fontsize=10)
        
        # Personnalisation
        ax.set_xlabel('Proportion (%)', fontsize=12, fontweight='bold')
        ax.set_ylabel('Types de prestations', fontsize=12, fontweight='bold')
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        
        # Définir les labels des axes Y (actes)
        ax.set_yticks(range(len(actes)))
        ax.set_yticklabels(actes, fontsize=10)
        
        # Inverser l'ordre pour avoir le plus gros en haut
        ax.invert_yaxis()
        
        # Grille pour la lisibilité
        ax.grid(axis='x', alpha=0.3, linestyle='--')
        ax.set_axisbelow(True)
        
        # Définir les limites de l'axe X
        max_prop = max(proportions)
        ax.set_xlim(0, max_prop * 1.15)
        
        plt.tight_layout()
        
        # Sauvegarde
        path = FIG_DIR / filename
        fig.savefig(path, dpi=300, bbox_inches='tight')
        plt.close(fig)
        return path
        
    except Exception as e:
        print(f"Erreur lors de la création du graphique {filename}: {e}")
        return None

# ------------------ Export Excel ------------------ #

def export_excel(global_indic: dict, reps: dict, evo: pd.DataFrame, comp_t1: pd.DataFrame, df: pd.DataFrame):
    # Préparation indicateurs avec labels FR
    indic_fr = {INDIC_LABELS.get(k, k): v for k, v in global_indic.items()}
    target = OUTPUT_EXCEL
    attempt = 0
    while True:
        try:
            writer_ctx = pd.ExcelWriter(target)
            break
        except PermissionError:
            attempt += 1
            if attempt > 1:
                # Générer un nouveau nom horodaté
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                target = f"rapport_prestations_complet_{ts}.xlsx"
            time.sleep(0.4)
    with writer_ctx as writer:
        (pd.Series(indic_fr).to_frame('Valeur')).to_excel(writer, sheet_name='Indicateurs')
        for key, val in reps.items():
            sheet = key.replace('repartition_', 'rep_')[:31]
            df_tmp = val.copy()
            if key != 'repartition_par_province' and not isinstance(df_tmp.index, pd.MultiIndex):
                df_tmp = df_tmp.reset_index()
                # Ajout de la colonne Numéro sauf pour le tableau statut
                if key != 'repartition_par_statut':
                    df_tmp.insert(0, 'Numéro', range(1, len(df_tmp) + 1))
                if len(df_tmp.columns):
                    first_col = df_tmp.columns[1] if 'Numéro' in df_tmp.columns else df_tmp.columns[0]
                    # Renommer première colonne de catégorie
                    df_tmp.rename(columns={first_col: HEADER_FIRST_MAP.get(key, first_col)}, inplace=True)
            elif key == 'repartition_par_province':
                # Ajouter Numéro en première colonne (sauf statut exclu déjà géré)
                if len(df_tmp):
                    df_tmp.insert(0, 'Numéro', range(1, len(df_tmp) + 1))
                    # Réordonner pour avoir Numéro, Province, Région puis métriques (Province avant Région comme dans HTML)
                    cols = df_tmp.columns.tolist()
                    order = ['Numéro']
                    if 'Province' in cols:
                        order.append('Province')
                    if 'Region' in cols:
                        order.append('Region')
                    for c in cols:
                        if c not in order:
                            order.append(c)
                    df_tmp = df_tmp[order]
            df_tmp.rename(columns=COL_LABELS, inplace=True)
            df_tmp.to_excel(writer, sheet_name=sheet, index=False)
        if not evo.empty:
            evo_fr = evo.rename(columns={'montant_total': 'Montant total', 'nb_prestations': 'Nombre de prestations'})
            evo_fr.to_excel(writer, sheet_name='Evolution_mensuelle')
        if not comp_t1.empty:
            comp_fr = comp_t1.rename(columns={'montant_total': 'Montant total', 'nb_prestations': 'Nombre de prestations'})
            comp_fr.to_excel(writer, sheet_name='Comparaison_T1')
        # Données brutes
        df.to_excel(writer, sheet_name='Donnees_nettoyees', index=False)
        # Doublons (reporting)
        if DUPLICATES_INFO:
            for key_sheet, df_dup in [
                ('Doublons_lignes', DUPLICATES_INFO.get('doublons_lignes')),
                ('Doublons_identifiant', DUPLICATES_INFO.get('doublons_identifiant')),
                ('Doublons_cle_comp', DUPLICATES_INFO.get('doublons_cle_composite'))
            ]:
                if isinstance(df_dup, pd.DataFrame) and not df_dup.empty:
                    # Ajouter une colonne Occurrences si pertinent (pour identifiant / clé composite)
                    if key_sheet != 'Doublons_lignes':
                        col_sub = []
                        if key_sheet == 'Doublons_identifiant':
                            col_sub = [c for c in df_dup.columns if 'identifiant' in c.lower()][:1]
                        elif key_sheet == 'Doublons_cle_comp':
                            col_sub = [c for c in ['adherent_code','date','acte','montant'] if c in df_dup.columns]
                        if col_sub:
                            occ = (df_dup.groupby(col_sub)
                                         .size()
                                         .reset_index(name='Occurrences')
                                         .sort_values('Occurrences', ascending=False))
                            occ.to_excel(writer, sheet_name=(key_sheet + '_occ')[:31], index=False)
                    df_dup.to_excel(writer, sheet_name=key_sheet[:31], index=False)
        # Lignes de dates corrigées
        if DATES_CORRIGEES_DF is not None and not DATES_CORRIGEES_DF.empty:
            tmp_dates = DATES_CORRIGEES_DF.copy()
            # Formatage lisible dates
            for dc_col in ['date_originale','date_corrigee']:
                if dc_col in tmp_dates.columns:
                    tmp_dates[dc_col] = tmp_dates[dc_col].dt.strftime('%d/%m/%Y')
            tmp_dates.to_excel(writer, sheet_name='Dates_corrigees', index=False)
    if target != OUTPUT_EXCEL:
        print(f"[INFO] Fichier Excel initial verrouillé. Export sauvegardé sous: {target}")

# ------------------ Export PDF ------------------ #

def export_pdf(global_indic: dict, reps: dict, images_paths, evo: pd.DataFrame, comp_t1: pd.DataFrame):
    if not _HAS_FPDF:
        print("Bibliothèque fpdf2 non installée -> export PDF ignoré.")
        return
    pdf = FPDF()
    # Appliquer protection si mots de passe définis
    try:
        if PDF_USER_PWD and PDF_OWNER_PWD and hasattr(pdf, 'set_protection'):
            pdf.set_protection(PDF_PERMISSIONS, user_pwd=PDF_USER_PWD, owner_pwd=PDF_OWNER_PWD)
        elif PDF_USER_PWD and PDF_OWNER_PWD and not hasattr(pdf, 'set_protection'):
            print("Info: méthode set_protection indisponible dans cette version de fpdf2 -> PDF non chiffré.")
    except Exception as e:
        print(f"Avertissement: échec application protection PDF ({e})")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    # Charger une police Unicode (DejaVuSans) si disponible, sinon fallback Helvetica
    unicode_font_path = None
    for candidate in [
        Path('DejaVuSans.ttf'),
        Path('C:/Windows/Fonts/DejaVuSans.ttf'),
        Path('C:/Windows/Fonts/DejaVuSansCondensed.ttf'),
        Path('C:/Windows/Fonts/arial.ttf')
    ]:
        if candidate.exists():
            unicode_font_path = candidate
            break
    if unicode_font_path:
        try:
            pdf.add_font('DejaVu', '', str(unicode_font_path), uni=True)
            pdf.add_font('DejaVu', 'B', str(unicode_font_path), uni=True)
            base_font = 'DejaVu'
        except Exception:
            base_font = 'Helvetica'
    else:
        base_font = 'Helvetica'
    pdf.set_font(base_font, 'B', 16)
    _sanitize = lambda s: str(s).replace('’', "'")
    pdf.cell(0, 10, _sanitize('Rapport Analytique des Prestations'), ln=1)
    # Logo (si présent) en haut à droite
    if LOGO_PATH.exists():
        try:
            pdf.image(str(LOGO_PATH), x=170, y=10, w=30)
        except Exception:
            pass
    pdf.set_font(base_font, '', 10)
    pdf.cell(0, 6, _sanitize(f"Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')}"), ln=1)
    # Synthèse introductive (résumé court des objectifs)
    pdf.set_font(base_font, 'B', 12)
    pdf.cell(0, 8, _sanitize("Introduction & objectifs (résumé)"), ln=1)
    pdf.set_font(base_font, '', 9)
    intro_resume = [
        _sanitize("Période : janv. 2024 – juil. 2025. Analyse des prestations remboursables nettoyées et fiabilisées."),
        _sanitize("Finalité : fournir une lecture stratégique des volumes, montants, évolutions et concentrations."),
        _sanitize("Axes : actes, types/sous types, centres, partenaires, zones géographiques, bénéficiaires & genre."),
        _sanitize("Temporalité : double lecture mensuelle (sensibilité) et trimestrielle (tendance)."),
        _sanitize("Usage : prioriser audits, prévention, renégociation et pilotage médico‑financier."),
    ]
    max_w = pdf.w - pdf.l_margin - pdf.r_margin
    for o in intro_resume:
        words = o.split()
        line = "- "
        for w in words:
            test = line + w + ' '
            if pdf.get_string_width(test) > max_w:
                pdf.cell(0, 5, line.rstrip(), ln=1)
                line = "  " + w + ' '
            else:
                line = test
        if line.strip():
            pdf.cell(0, 5, line.rstrip(), ln=1)
    pdf.ln(2)
    # Indicateurs globaux
    pdf.set_font(base_font, 'B', 12)
    pdf.cell(0, 8, '1. Indicateurs globaux', ln=1)
    pdf.set_font(base_font, '', 9)
    for k, v in global_indic.items():
        label = INDIC_LABELS.get(k, k)
        if isinstance(v, (int, float)) and not isinstance(v, bool):
            if 'taux' in k:
                val_str = f"{v:.1f}%"
            else:
                val_str = f"{v:,.0f}".replace(',', ' ')
        elif isinstance(v, pd.Timestamp):
            val_str = v.strftime('%d/%m/%Y') if not pd.isna(v) else ''
        else:
            val_str = str(v)
        pdf.cell(95, 6, label, border=1)
        pdf.cell(95, 6, val_str, border=1, ln=1)
    # Répartitions clefs
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 8, '2. Répartitions', ln=1)
    pdf.set_font('Helvetica', '', 9)
    header_first_map = {
        'repartition_par_acte': 'Acte',
        'repartition_par_centre': 'Centre de santé',
        'repartition_par_region': 'Région',
        'repartition_par_province': 'Province',  # géré séparément (Province + Région)
        'repartition_par_type': 'Type',
        'repartition_par_sous_type': 'Sous-type',
        'repartition_par_statut': 'Statut'
    }
    for key in ['repartition_par_acte', 'repartition_par_centre', 'repartition_par_region', 'repartition_par_province', 'repartition_par_type', 'repartition_par_sous_type', 'repartition_par_statut', 'repartition_par_beneficiaire', 'repartition_par_genre']:
        if key in reps:
            titre = key.replace('repartition_par_', 'Répartition par ').replace('_', ' ').title()
            pdf.set_font(base_font, 'B', 10)
            pdf.cell(0, 6, titre, ln=1)
            pdf.set_font(base_font, '', 8)
            df_source = reps[key]
            # Pour province avec Region colonne explicite
            if key == 'repartition_par_province' and 'Region' in df_source.columns:
                df_rep = df_source.copy()
            else:
                df_rep = df_source.rename(columns=COL_LABELS)
            df_rep = df_rep.head(15)
            cols = df_rep.columns.tolist()
            col_widths = []
            # Largeur dynamique: catégorie 60 puis repartir  (max 3 colonnes esperées)
            first_col_width = 40 if key == 'repartition_par_province' else 60
            remaining = 190 - first_col_width
            if len(cols):
                per = remaining / len(cols)
                col_widths = [per for _ in cols]
            # En-têtes
            pdf.set_fill_color(230, 230, 230)
            label_first = header_first_map.get(key, 'Catégorie')
            if key == 'repartition_par_province':
                # Colonnes: N° | Province | Région (fusionnée) | métriques
                num_col_w = 12
                pdf.cell(num_col_w, 6, 'N°', border=1, fill=True)
                pdf.cell(first_col_width, 6, 'Province', border=1, fill=True)
                pdf.cell(first_col_width, 6, 'Région', border=1, fill=True)
                metric_cols = [c for c in cols if c not in ('Region','Province')]
                width_metrics_total = 190 - first_col_width - first_col_width - num_col_w
                for col in metric_cols:
                    pdf.cell( width_metrics_total / max(1,len(metric_cols)), 6, COL_LABELS.get(col,col), border=1, fill=True)
                pdf.ln(6)
                # Lignes avec fusion logique (répéter région seulement si change)
                last_region = None
                numero = 1
                for _, row in df_rep.iterrows():
                    reg = row.get('Region','')
                    prov = row.get('Province','')
                    # Numéro + Province toujours
                    pdf.cell(num_col_w, 6, str(numero), border=1)
                    pdf.cell(first_col_width, 6, str(prov)[:18], border=1)
                    # Région fusionnée visuellement (cellule vide si répétée)
                    if reg != last_region:
                        pdf.cell(first_col_width, 6, str(reg)[:18], border=1)
                        last_region = reg
                    else:
                        pdf.cell(first_col_width, 6, '', border=1)
                    for col in metric_cols:
                        val = row[col]
                        if isinstance(val, (int,float)) and not isinstance(val,bool):
                            if float(val).is_integer():
                                val_str = f"{int(val):,}".replace(',', ' ')
                            else:
                                val_str = f"{val:,.2f}".replace(',', ' ').replace('.', ',')
                        else:
                            val_str = str(val)
                        pdf.cell( width_metrics_total / max(1,len(metric_cols)), 6, val_str, border=1)
                    pdf.ln(6)
                    numero += 1
                pdf.ln(4)
                continue
            pdf.cell(first_col_width, 6, label_first, border=1, fill=True)
            for w, col in zip(col_widths, cols):
                pdf.cell(w, 6, col, border=1, fill=True)
            pdf.ln(6)
            for idx, row in df_rep.iterrows():
                pdf.cell(first_col_width, 6, str(idx)[:30], border=1)
                for w, col in zip(col_widths, cols):
                    val = row[col]
                    if isinstance(val, (int, float)) and not isinstance(val, bool):
                        val_str = f"{val:,.0f}".replace(',', ' ')
                    else:
                        val_str = str(val)
                    pdf.cell(w, 6, val_str, border=1)
                pdf.ln(6)
            pdf.ln(4)
    # Images
    for img in images_paths:
        pdf.add_page()
        pdf.set_font(base_font, 'B', 11)
        pdf.cell(0, 6, img.stem.replace('_', ' ').title(), ln=1)
        try:
            pdf.image(str(img), w=180)
        except Exception as e:
            pdf.set_font(base_font, '', 9)
            pdf.multi_cell(0, 5, f"Image introuvable: {e}")
    # Page explications & notation
    pdf.add_page()
    pdf.set_font(base_font, 'B', 12)
    pdf.cell(0, 8, _sanitize('4. Notes & Explications'), ln=1)
    pdf.set_font(base_font, '', 9)
    notes = [
        "Notation 1 877 340 (124) : le premier nombre est le montant total (espaces = séparateurs de milliers), le nombre entre parenthèses est le nombre de prestations.",
        "Abréviations des montants sur certaines courbes : 12,3 k = 12 300 ; 1,45 M = 1 450 000 ; 2,10 Md = 2 100 000 000.",
    "Montant total par acte : compare le coût cumulé par type d'acte (consultations, examens, imagerie, pharmacie, etc.) pour identifier les postes majeurs.",
    "Répartition des montants par acte (camembert) : part relative de chaque catégorie d'actes dans la dépense totale.",
    "Top 10 centres (montant / nombre) : centres présentant les montants remboursés les plus élevés et/ou le plus grand nombre de prestations.",
    "Top 10 partenaires : partenaires avec les montants remboursés les plus élevés et/ou le plus grand nombre de prestations.",
        "Répartition par statut de traitement : proportion de prestations acceptées vs autres statuts.",
        "Répartition par région / province : concentration géographique des montants et volumes.",
        "Évolution mensuelle : barres = nombre de prestations, ligne = montant total, annotations = montant abrégé + (nombre).",
        "Évolution trimestrielle : tendance consolidée par trimestre (montant + nombre).",
        "Top 10 adhérents (montant) : adhérents dont le total des prestations est le plus élevé (indicateur de concentration du risque)."
    ]
    for n in notes:
        # simple wrapping
        max_w = pdf.w - pdf.l_margin - pdf.r_margin
        words = n.split()
        line = "- "
        for w in words:
            test = line + w + ' '
            if pdf.get_string_width(test) > max_w:
                pdf.cell(0, 5, _sanitize(line.rstrip()), ln=1)
                line = "  " + w + ' '
            else:
                line = test
        if line.strip():
            pdf.cell(0, 5, _sanitize(line.rstrip()), ln=1)
    # Conclusion
    pdf.ln(3)
    pdf.set_font(base_font, 'B', 12)
    pdf.cell(0, 8, _sanitize('5. Conclusion'), ln=1)
    pdf.set_font(base_font, '', 9)
    conclusion_text = (
        "Cette analyse a permis de transformer des données brutes en une lecture stratégique, mettant en lumière les principaux contributeurs aux dépenses, "
        "les déséquilibres potentiels et les tendances d’évolution. Les constats dégagés ouvrent la voie à des actions concrètes : renforcer la maîtrise médico-financière, cibler les audits, "
        "optimiser les partenariats et anticiper les risques. L’enjeu est désormais de traduire ces enseignements en décisions opérationnelles afin de soutenir la soutenabilité du régime et d’outiller la gouvernance dans son pilotage stratégique." )
    # Normalisation de certains caractères pouvant poser problème de largeur / encodage
    conclusion_text = (conclusion_text
                       .replace('\u00A0', ' ')
                       .replace('\u202F', ' ')
                       .replace('\u2011', '-')  # trait d'union insécable
                       .replace('\u2013', '-')  # tiret demi-cadrat
                       .replace('\u2014', '-')  # tiret cadrat
                      )
    # Wrap conclusion (algorithme manuel similaire aux puces pour éviter l'exception fpdf"Not enough horizontal space")
    max_w = pdf.w - pdf.l_margin - pdf.r_margin
    words = conclusion_text.split()
    line = ''
    for w in words:
        candidate = (line + ' ' + w).strip()
        if pdf.get_string_width(_sanitize(candidate)) > max_w and line:
            pdf.cell(0, 5, _sanitize(line.rstrip()), ln=1)
            line = w
        else:
            line = candidate
    if line.strip():
        pdf.cell(0, 5, _sanitize(line.rstrip()), ln=1)
    pdf.output(OUTPUT_PDF)
    print(f"Export PDF: {OUTPUT_PDF}")

# ------------------ Main ------------------ #

def main():
    df = charger(SOURCE_FILE)
    global_indic = compute_global(df)
    reps = repartitions(df)
    evo, comp_t1 = analyse_temporelle(df)

    # Graphiques principaux
    montant_col = next((c for c in df.columns if c.lower().startswith('montant')), None)
    images_paths = []
    if 'repartition_par_acte' in reps:
        acte_sum = reps['repartition_par_acte']['sum']
        acte_count = reps['repartition_par_acte']['count'] if 'count' in reps['repartition_par_acte'].columns else None
        p1 = plot_bar(acte_sum, 'Montant total par acte', 'montant_par_acte.png', counts_series=acte_count)
        if p1: images_paths.append(p1)
        p1b = plot_pie_group_small(reps['repartition_par_acte']['sum'], 'Répartition des montants par acte', 'montant_par_acte_pie.png')
        if p1b: images_paths.append(p1b)
    if 'repartition_par_centre' in reps:
        centre_sum = reps['repartition_par_centre']['sum'].head(10)
        centre_count = reps['repartition_par_centre']['count'] if 'count' in reps['repartition_par_centre'].columns else None
        if centre_count is not None:
            centre_count = centre_count.head(10)
        p2 = plot_bar(centre_sum, 'Top 10 centres (montant)', 'top10_centres.png', horizontal=True, counts_series=centre_count)
        if p2: images_paths.append(p2)
        # Trier par nombre de prestations en ordre décroissant pour le graphique
        centre_count_sorted = reps['repartition_par_centre']['count'].sort_values(ascending=False).head(10)
        p2b = plot_bar(centre_count_sorted, 'Top 10 centres (nombre prestations)', 'top10_centres_nbp.png', horizontal=True)
        if p2b: images_paths.append(p2b)
    if 'repartition_par_partenaire' in reps:
        part_sum = reps['repartition_par_partenaire']['sum'].head(10)
        part_count = reps['repartition_par_partenaire']['count'] if 'count' in reps['repartition_par_partenaire'].columns else None
        if part_count is not None:
            part_count = part_count.head(10)
        pp1 = plot_bar(part_sum, 'Top 10 partenaires (montant)', 'top10_partenaires.png', horizontal=True, counts_series=part_count)
        if pp1: images_paths.append(pp1)
        pp2 = plot_bar(reps['repartition_par_partenaire']['count'].head(10), 'Top 10 partenaires (nombre prestations)', 'top10_partenaires_nbp.png', horizontal=True)
        if pp2: images_paths.append(pp2)
    # Top 10 adhérents (montant total)
    montant_col = next((c for c in df.columns if c.lower().startswith('montant')), None)
    adherent_col = next((c for c in df.columns if 'adherent_code' in c.lower()), None)
    if montant_col and adherent_col and adherent_col in df.columns:
        grp_ad = df.groupby(adherent_col)
        top_adherents = grp_ad[montant_col].sum().sort_values(ascending=False).head(10)
        counts_adherents = grp_ad[montant_col].count().reindex(top_adherents.index)
        pa = plot_bar(top_adherents, 'Top 10 adhérents (montant)', 'top10_adherents_montant.png', horizontal=True, counts_series=counts_adherents)
        if pa: images_paths.append(pa)
    if 'repartition_par_statut' in reps:
        ps = plot_pie_group_small(reps['repartition_par_statut']['count'], 'Répartition des prestations par statut', 'repartition_statut_pie.png', threshold=0.01)
        if ps: images_paths.append(ps)
    # Histogrammes Type / Sous-type (verticaux pour lisibilité et cohérence)
    if 'repartition_par_type' in reps:
        try:
            type_sum = reps['repartition_par_type']['sum']
            g_type = plot_bar(type_sum, 'Montant total par type', 'montant_par_type.png', annotate_vertical=True)
            if g_type: images_paths.append(g_type)
        except Exception:
            pass
    if 'repartition_par_sous_type' in reps:
        try:
            st_sum = reps['repartition_par_sous_type']['sum'].head(30)  # élargir à top 30 si disponible
            g_soustype = plot_bar(st_sum, 'Montant total par sous-type (Top 30)', 'montant_par_sous_type.png')
            if g_soustype: images_paths.append(g_soustype)
        except Exception:
            pass
    # Répartition par bénéficiaire (camembert)
    if 'repartition_par_beneficiaire' in reps:
        try:
            pb = plot_pie_group_small(reps['repartition_par_beneficiaire']['count'], 'Répartition des prestations par bénéficiaire', 'repartition_beneficiaire_pie.png', threshold=0.01)
            if pb: images_paths.append(pb)
        except Exception:
            pass
    # Répartition par genre (camembert)
    if 'repartition_par_genre' in reps:
        try:
            pg = plot_pie_group_small(reps['repartition_par_genre']['count'], 'Répartition des prestations par genre', 'repartition_genre_pie.png', threshold=0.01)
            if pg: images_paths.append(pg)
        except Exception:
            pass
    # Graphiques région & province
    reg_path = plot_region_distribution(reps)
    if reg_path: images_paths.append(reg_path)
    prov_path = plot_province_distribution(reps)
    if prov_path: images_paths.append(prov_path)
    if not evo.empty and montant_col:
        # Préparation données (sans moyenne mobile)
        evo_ma = evo.copy()
        x_vals = list(range(len(evo_ma.index)))
        fig, ax1 = plt.subplots(figsize=(13,5))
        # Barres pour nb prestations
        ax1.bar(x_vals, evo_ma['nb_prestations'], color='#b0c4de', alpha=0.65, label='Nombre de prestations')
        ax1.set_ylabel('Nombre de prestations')
        ax1.set_xlabel('Mois')
        ax2 = ax1.twinx()
        ax2.plot(x_vals, evo_ma['montant_total'], marker='o', color='#d35400', label='Montant total')
        ax2.set_ylabel('Montant (FR)')
        ax2.set_title('Évolution mensuelle – Montants & Prestations')
        # Légendes combinées
        lines, labels = [], []
        for ax in [ax1, ax2]:
            l, lab = ax.get_legend_handles_labels()
            lines.extend(l); labels.extend(lab)
        ax2.legend(lines, labels, loc='upper left')
        # Annotations abrégées montants
        amt_texts = [f"{format_abbrev(v)}\n({int(n)})" for v, n in zip(evo_ma['montant_total'], evo_ma['nb_prestations'])]
        annotate_line_no_overlap(ax2, x_vals, list(evo_ma['montant_total']), amt_texts, min_gap_frac=0.025, base_x_offset=0.4)
        ax1.set_xticks(x_vals)
        ax1.set_xticklabels(evo_ma.index, rotation=45, ha='right')
        ax2.grid(axis='y', linestyle=':', alpha=0.9, linewidth=0.1)
        fig.tight_layout()
        # Bordure carrée figure
        rect = Rectangle((0.005,0.005),0.99,0.99, transform=fig.transFigure, fill=False, lw=1.4, edgecolor='#333')
        fig.patches.append(rect)
        evol_path = FIG_DIR / 'evolution_mensuelle.png'
        fig.savefig(evol_path)
        plt.close(fig)
        images_paths.append(evol_path)
        # Évolution trimestrielle
        date_col = next((c for c in df.columns if c.lower() == 'date'), None)
        if date_col:
            temp_q = df[[date_col, montant_col]].copy()
            temp_q['trimestre'] = temp_q[date_col].dt.to_period('Q').astype(str)
            q_evo = temp_q.groupby('trimestre')[montant_col].agg(['sum','count']).rename(columns={'sum':'montant_total','count':'nb_prestations'})
            if not q_evo.empty:
                fig2, ax2 = plt.subplots(figsize=(9,5))
                sns.lineplot(x=q_evo.index, y=q_evo['montant_total'], marker='o', ax=ax2)
                ax2.set_title('Évolution trimestrielle des montants')
                ax2.set_xlabel('Trimestre')
                ax2.set_ylabel('Montant total')
                for x, (y, n) in zip(q_evo.index, zip(q_evo['montant_total'], q_evo['nb_prestations'])):
                    txt = f"{int(y):,} ({int(n)})".replace(',', ' ')
                    ax2.annotate(txt, (x, y), textcoords='offset points', xytext=(0,7), ha='center', fontsize=8)
                plt.xticks(rotation=0)
                plt.tight_layout()
                # Bordure carrée figure
                rect2 = Rectangle((0.005,0.005),0.99,0.99, transform=fig2.transFigure, fill=False, lw=1.4, edgecolor='#333')
                fig2.patches.append(rect2)
                evol_q_path = FIG_DIR / 'evolution_trimestrielle.png'
                fig2.savefig(evol_q_path)
                plt.close(fig2)
                images_paths.append(evol_q_path)

    export_excel(global_indic, reps, evo, comp_t1, df)
    try:
        print('Images générées:', [p.stem for p in images_paths])
    except Exception:
        pass
    generer_rapport_html(global_indic, reps, evo, comp_t1, images_paths)
    export_pdf(global_indic, reps, images_paths, evo, comp_t1)
    export_word(global_indic, reps, evo, comp_t1, images_paths)

    print('--- Résumé ---')
    for k, v in global_indic.items():
        print(f"{INDIC_LABELS.get(k,k)}: {v}")
    print(f"Export Excel: {OUTPUT_EXCEL}")
    print(f"Graphiques dans: {FIG_DIR}/")
    print("Rapport HTML: rapport_prestations.html")
    if 'top10_centres' in reps:
        print('Top centres (montant):')
        print(reps['top10_centres'].head(5))

def _img_tag_from_path(path: Path) -> str:
    try:
        with open(path, 'rb') as f:
            b64 = base64.b64encode(f.read()).decode('utf-8')
        return f'<img src="data:image/png;base64,{b64}" alt="{path.stem}" style="max-width:100%;height:auto;" />'
    except Exception as e:
        return f'<p>Image {path} indisponible ({e})</p>'

def export_word(global_indic: dict, reps: dict, evo: pd.DataFrame, comp_t1: pd.DataFrame, images_paths):
    """Export Word en répliquant l'ordre et la structure du HTML (sections, titres, tableaux, figures)."""
    if not _HAS_DOCX:
        print("python-docx non installé -> export Word ignoré.")
        return
    # Re-générer le HTML ou réutiliser celui déjà stocké
    global LAST_HTML_REPORT
    if LAST_HTML_REPORT is None:
        generer_rapport_html(global_indic, reps, evo, comp_t1, images_paths)
    html_text = LAST_HTML_REPORT or ''
    try:
        # Tentative BeautifulSoup puis fallback lxml
        parser_used = 'bs4'
        try:
            from bs4 import BeautifulSoup  # type: ignore
            soup = BeautifulSoup(html_text, 'html.parser')
        except Exception:
            import lxml.html  # type: ignore
            parser_used = 'lxml'
            soup = lxml.html.fromstring(html_text)
        doc = Document()
        # Ajout titre global si présent
        def plain(text: str):
            return text.replace('\u00a0',' ').strip()
        # Fonctions utilitaires (versions bs4 / lxml unifiées via attributs)
        def iter_body_children():
            if parser_used == 'bs4':
                for child in soup.body.children:
                    if getattr(child, 'name', None):
                        yield child
            else:
                body = soup.find('body') if hasattr(soup,'find') else None
                body = body or soup
                for child in body.getchildren():
                    yield child
        def get_name(el):
            return getattr(el,'name', None) if parser_used=='bs4' else el.tag
        def get_text(el):
            return el.get_text(strip=True) if parser_used=='bs4' else ''.join(el.itertext()).strip()
        def children(el):
            if parser_used=='bs4':
                for c in el.children:
                    if getattr(c,'name',None):
                        yield c
            else:
                for c in el.getchildren():
                    yield c
        def find_all(el, tags):
            if parser_used=='bs4':
                return el.find_all(tags, recursive=False)
            else:
                return [c for c in el.getchildren() if c.tag in tags]
        def add_heading(txt, level):
            if not txt: return
            if level==0:
                doc.add_heading(txt, 0)
            else:
                doc.add_heading(txt, level=level)
        def render_table(el):
            rows = []
            if parser_used=='bs4':
                for r in el.find_all('tr'):
                    rows.append([plain(c.get_text()) for c in r.find_all(['th','td'])])
            else:
                for r in el.findall('.//tr'):
                    cells = []
                    for c in r.findall('./th') + r.findall('./td'):
                        cells.append(plain(''.join(c.itertext())))
                    if cells:
                        rows.append(cells)
            if not rows: return
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=1, cols=ncols)
            # fonction de prettify pour rendre les clés lisibles
            def prettify_label(s: str) -> str:
                if not s or not isinstance(s, str):
                    return str(s)
                # alias map pour cas fréquents
                aliases = {
                    'analyse_biomedicale': 'Analyse biomédicale',
                    'dentaire_auditif': 'Dentaire & Auditif',
                    'maternite': 'Maternité',
                    'hospitalisation': 'Hospitalisation',
                    'optique': 'Optique',
                    'consultation': 'Consultation',
                    'pharmacie': 'Pharmacie',
                    'autre': 'Autre'
                }
                key = s.strip()
                low = key.lower()
                if low in aliases:
                    return aliases[low]
                # Si ressemble à une clé (underscores, pas d'espaces), remplacer underscores et capitaliser
                if '_' in key or (key == key.lower() and ' ' not in key and len(key) < 40):
                    parts = key.replace('_', ' ').split()
                    return ' '.join(p.capitalize() for p in parts)
                return key
            for j, val in enumerate(rows[0]):
                tbl.rows[0].cells[j].text = prettify_label(val)
            for r in rows[1:]:
                row = tbl.add_row().cells
                for j, val in enumerate(r):
                    # appliquer prettify sur la première colonne si nécessaire
                    if j == 0:
                        row[j].text = prettify_label(val)
                    else:
                        row[j].text = val
        import base64, tempfile
        def render_img(el):
            src = el.get('src') if parser_used!='bs4' else el.get('src','')
            if not src: return
            try:
                if src.startswith('data:image'):
                    b64 = src.split(',',1)[1]
                    data = base64.b64decode(b64)
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                        tmp.write(data)
                        tmp.flush()
                        doc.add_picture(tmp.name, width=Inches(5.5))
                elif Path(src).exists():
                    doc.add_picture(src, width=Inches(5.5))
            except Exception:
                pass
        def render_figure(fig):
            # image
            if parser_used=='bs4':
                img = fig.find('img')
                if img:
                    render_img(img)
                cap = fig.find('figcaption')
                if cap:
                    p = doc.add_paragraph(plain(cap.get_text()))
                    p.italic = True
            else:
                img = fig.find('.//img')
                if img is not None:
                    render_img(img)
                cap = fig.find('.//figcaption')
                if cap is not None:
                    p = doc.add_paragraph(plain(''.join(cap.itertext())))
                    p.italic = True
        # Parcours
        for node in iter_body_children():
            name = get_name(node)
            if name in {'h1','h2','h3'}:
                lvl = {'h1':0,'h2':1,'h3':2}[name]
                add_heading(get_text(node), lvl)
            elif name == 'p':  # paragraphes de premier niveau (ex: explication indicateurs globaux)
                txt = get_text(node)
                if txt:
                    doc.add_paragraph(txt)
            elif name == 'table':  # tableau indicateurs globaux
                render_table(node)
            elif name == 'figure':
                render_figure(node)
            elif name == 'section':
                # parcourir contenu dans l'ordre
                for ch in children(node):
                    cname = get_name(ch)
                    if cname in {'h1','h2','h3'}:
                        add_heading(get_text(ch), {'h1':0,'h2':1,'h3':2}[cname])
                    elif cname == 'p':
                        txt = get_text(ch)
                        if txt: doc.add_paragraph(txt)
                    elif cname in {'ul','ol'}:
                        li_tags = find_all(ch, ['li']) if parser_used!='bs4' else ch.find_all('li', recursive=False)
                        for li in li_tags:
                            style = 'List Number' if cname=='ol' else 'List Bullet'
                            doc.add_paragraph(get_text(li), style=style)
                    elif cname == 'table':
                        # Vérifier si c'est le tableau des centres à masquer dans Word
                        table_class = node.get('class', []) if hasattr(node, 'get') else []
                        if isinstance(table_class, str):
                            table_class = table_class.split()
                        skip_centres_table = 'centres-table-skip-word' in table_class
                        if skip_centres_table:
                            # Ne pas insérer ce tableau dans Word - juste mettre un commentaire
                            doc.add_paragraph('[Tableau "classement des centres" temporairement masqué dans ce document]', style=None)
                        else:
                            render_table(ch)
                    elif cname == 'figure':
                        render_figure(ch)
                    else:
                        # scan sous-éléments potentiels table/figure
                        if cname == 'div':
                            for sub in children(ch):
                                sname = get_name(sub)
                                if sname == 'table':
                                    # Vérifier si c'est le tableau des centres à masquer dans Word
                                    table_class = sub.get('class', []) if hasattr(sub, 'get') else []
                                    if isinstance(table_class, str):
                                        table_class = table_class.split()
                                    skip_centres_table = 'centres-table-skip-word' in table_class
                                    if skip_centres_table:
                                        # Ne pas insérer ce tableau dans Word - juste mettre un commentaire
                                        doc.add_paragraph('[Tableau "classement des centres" temporairement masqué dans ce document]', style=None)
                                    else:
                                        render_table(sub)
                                elif sname == 'figure':
                                    render_figure(sub)
        # Sauvegarde avec fallback si verrou
        try:
            doc.save(OUTPUT_DOCX)
        except PermissionError:
            alt = OUTPUT_DOCX.replace('.docx', f"_{datetime.now().strftime('%H%M%S')}.docx")
            doc.save(alt)
            print(f"Export Word verrouillé, sauvegarde alternative: {alt}")
        else:
            print(f"Export Word: {OUTPUT_DOCX}")
    except Exception as e:
        print(f"Erreur export Word: {e}")

def generer_rapport_html(global_indic: dict, reps: dict, evo: pd.DataFrame, comp_t1: pd.DataFrame, images_paths):
    def fmt(v):
        if isinstance(v, (int, float)) and not isinstance(v, bool):
            if pd.isna(v):
                return ''
            # entier ?
            if float(v).is_integer():
                return f"{int(v):,}".replace(',', ' ')
            # décimal 2 chiffres, séparateur milliers espace, virgule française
            s = f"{v:,.2f}"  # 1,234.56
            s = s.replace(',', ' ').replace('.', ',')
            return s
        return str(v)
    # Exclure des indicateurs globaux : montants payés / non payés (ils vont dans Traitement des données)
    exclude_keys = {'montant_paye_total','montant_non_paye','pourcentage_paye_pct'}
    rows_indic = '\n'.join(
        f'<tr><td>{INDIC_LABELS.get(k,k)}</td><td>{fmt(v if not isinstance(v, pd.Timestamp) else v.strftime("%d/%m/%Y"))}</td></tr>'
        for k, v in global_indic.items() if k not in exclude_keys)
    # Ajout synthèse doublons (compteurs) dans indicateurs globaux
    if DUPLICATES_INFO:
        synth_dups = [
            ('Doublons (lignes complètes)', DUPLICATES_INFO.get('nb_doublons_lignes', 0)),
            ("Doublons (identifiant)", DUPLICATES_INFO.get('nb_doublons_identifiant', 0)),
            ("Doublons (clé composite)", DUPLICATES_INFO.get('nb_doublons_cle_composite', 0)),
        ]
        for label, val in synth_dups:
            rows_indic += f"<tr><td>{label}</td><td>{val}</td></tr>"
    sections = []  # legacy list (will be replaced by unified blocks)
    table_htmls: dict[str, dict] = {}
    # Descriptions des tableaux (inline)
    table_desc = {
        'repartition_par_acte': (
            "Ce tableau offre une lecture structurée des actes médicaux en rapprochant simultanément le poids financier et la fréquence. "
            "L'objectif est de dégager un noyau d'actes dominants qui concentrent l'effort budgétaire. Une catégorie très coûteuse mais peu fréquente appelle des vérifications qualitatives (pertinence clinique, tarification), tandis qu'une catégorie très fréquente et modérément coûteuse peut devenir prioritaire pour des actions préventives ou de rationalisation. "
            "Cette double perspective évite de se laisser guider par le seul montant agrégé et prépare les arbitrages ciblés."
        ),
        'repartition_par_centre': (
            "Cette répartition éclaire la contribution respective des centres de santé. Elle distingue les sites qui concentrent les montants de ceux qui concentrent l'activité brute. "
            "Un centre haut en montant mais moyen en volume peut refléter une spécialisation coûteuse (imagerie, hospitalisation complexe) alors qu'un centre élevé en volume mais modeste en montant relève souvent de soins courants nécessitant fluidité et qualité plutôt que négociation tarifaire. "
            "L'analyse permet de prioriser les visites terrain, audits, ou renégociations contractuelles."
        ),
        'repartition_par_region': (
            "La vision régionale sert de filtre macro pour détecter des asymétries territoriales : surreprésentation des dépenses, sous‑utilisation relative ou polarisation géographique. "
            "Elle permet d'anticiper des tensions locales (saturation, dérives de pratiques) et d'interroger l'équité d'accès. Une région sous‑pondérée en volume mais avec un coût moyen élevé peut indiquer des actes plus lourds ou une offre moins diversifiée. "
            "Ce diagnostic de premier niveau guide l'examen provincial fin."
        ),
        'repartition_par_province': (
            "Le passage à l'échelle provinciale précise l'origine réelle des écarts détectés plus haut. "
            "On vérifie si une région dominante repose sur quelques provinces moteurs ou sur un socle homogène. Cette granularité facilite la mise en place d'actions correctives localisées (sensibilisation, renforcement du réseau, contrôle ciblé). "
            "Elle évite aussi les généralisations hâtives issues d'agrégats régionaux."
        ),
        'repartition_par_type': (
            "La catégorisation par grands types agrège suffisamment pour révéler la structure fonctionnelle globale de la dépense sans diluer les signaux forts. "
            "L'histogramme des montants par type permet de visualiser immédiatement les domaines dominants et de confirmer les priorités d'analyse. "
            "Ce niveau intermédiaire aide à prioriser : sur quels domaines (externe, pharmacie, hospitalisation…) concentrer les efforts d'analyse détaillée ? "
            "Il constitue un pivot entre la vision synthétique et le grain fin des sous‑types."
        ),
        'repartition_par_sous_type': (
            "Les sous‑types dévoilent les poches précises de dépense ou de dynamisme. "
            "L'histogramme (Top 15 sous‑types) met en évidence les niches coûteuses masquées au niveau supérieur et oriente les investigations ciblées. "
            "Ils permettent d'identifier des segments candidats à des protocoles de soins, à une revue tarifaire ou à une action de prévention. Une sous‑catégorie en forte progression relative mérite une surveillance anticipative avant qu'elle ne pèse lourdement sur l'enveloppe. "
            "Cette granularité est essentielle pour transformer l'observation en plan d'action."
        ),
        'repartition_par_partenaire': (
            "Le tableau des partenaires met en regard le poids financier et le volume traité par chaque entité contractée. "
            "Il permet d'identifier les relations critiques (forte part du portefeuille), de surveiller les dépendances excessives et de préparer des discussions ciblées sur les pratiques de facturation ou les engagements de service. "
            "Une concentration élevée justifie un plan de sécurisation (clauses, diversification progressive) tandis qu'une dispersion peut indiquer un réseau équilibré mais plus complexe à piloter."
        ),
        'repartition_par_statut': (
            "La distribution des statuts évalue la performance opérationnelle du flux de traitement. "
            "Un taux de rejet ou d'état intermédiaire élevé peut résulter d'erreurs de saisie, de pièces justificatives incomplètes ou d'ambiguïtés procédurales. "
            "Suivre son évolution dans le temps fournit un indicateur indirect d'efficacité interne et de maturité du réseau de prestataires."
        ),
        'evolution_mensuelle_tableau': (
            "Le détail mensuel constitue la base de vérification et de recalcul. "
            "Il permet de confronter les tendances visuelles à la matérialité des chiffres, de repérer des ruptures (sauts de niveau, creux saisonniers) et de préparer d'éventuelles projections simples (extrapolation linéaire, moyenne glissante). "
            "Sa structure textuelle (mois en toutes lettres) facilite la communication non technique."
        ),
        'evolution_trimestrielle_tableau': (
            "La consolidation trimestrielle filtre le bruit court terme pour révéler les inflexions structurelles. "
            "Elle constitue un format adapté aux instances de gouvernance qui privilégient la trajectoire globale (accélération, plateau, repli). "
            "Ce niveau est aussi pertinent pour aligner budgétairement les anticipations avec la réalité observée."
        )
    }
    # Ajouts descriptions spécifiques bénéficiaire / genre
    table_desc['repartition_par_beneficiaire'] = (
        "Cette ventilation distingue les catégories de bénéficiaires (adhérent principal, ayant droit, etc.) en rapprochant charge financière et fréquence. "
        "Elle permet de vérifier si une catégorie spécifique pèse de manière disproportionnée et d'orienter des actions de prévention ou d'accompagnement ciblé."
    )
    table_desc['repartition_par_genre'] = (
        "La distribution par genre met en évidence d'éventuelles asymétries de recours. "
        "Elle constitue un indicateur d'équité et peut déclencher des analyses complémentaires (accès, comportements de soins, pathologies différenciées)."
    )
    # Tableaux principaux
    header_first_map = {
        'repartition_par_acte': 'Acte',
        'repartition_par_centre': 'Centre de santé',
        'repartition_par_region': 'Région',
        'repartition_par_province': 'Province',  # Province + Région handled séparément
        'repartition_par_type': 'Type',
        'repartition_par_sous_type': 'Sous-type',
        'repartition_par_statut': 'Statut'
    }
    for key in ['repartition_par_acte','repartition_par_centre','repartition_par_partenaire','repartition_par_region','repartition_par_province','repartition_par_type','repartition_par_sous_type','repartition_par_statut','repartition_par_beneficiaire','repartition_par_genre']:
        if key in reps:
            tbl = reps[key].copy()
            tbl_fmt = tbl.copy()
            for col in tbl_fmt.columns:
                if col in {'sum','count','mutualistes_distincts'}:
                    tbl_fmt[col] = tbl_fmt[col].apply(lambda x: fmt(x))
                elif col == 'mean':
                    tbl_fmt[col] = tbl_fmt[col].apply(lambda x: fmt(x))
            tbl_fmt.rename(columns=COL_LABELS, inplace=True)
            original_len = len(tbl_fmt)
            if key == 'repartition_par_province' and 'Region' in tbl.columns:
                # Construire HTML avec fusion région (rowspan) + formatage numérique
                dfp = tbl.copy()
                metric_cols = [c for c in dfp.columns if c not in ['Region','Province']]
                for mc in metric_cols:
                    dfp[mc] = dfp[mc].apply(lambda x: fmt(x))
                rows_html = []
                last_region = None
                region_counts = dfp['Region'].value_counts()
                region_rowspans = region_counts.to_dict()
                num = 1
                for _, r in dfp.iterrows():
                    region = r['Region']
                    province = r['Province']
                    row_cells = [f"<td>{num}</td>", f"<td>{province}</td>"]
                    if region != last_region:
                        rowspan = region_rowspans.get(region,1)
                        row_cells.append(f"<td rowspan='{rowspan}' style='font-weight:600;background:#f9f9f9;'>{region}</td>")
                        last_region = region
                    for met_col in metric_cols:
                        row_cells.append(f"<td>{r[met_col]}</td>")
                    rows_html.append('<tr>' + ''.join(row_cells) + '</tr>')
                    num += 1
                header = '<tr><th>Numéro</th><th>Province</th><th>Région</th>' + ''.join(f'<th>{COL_LABELS.get(c,c)}</th>' for c in metric_cols) + '</tr>'
                html_table_inner = f"<table class='tbl'><thead>{header}</thead><tbody>{''.join(rows_html)}</tbody></table>"
                html_table = f"<div class='scroll-box'>{html_table_inner}</div>" if original_len>20 else html_table_inner
                titre_tbl = key.replace('repartition_par_', 'Répartition par ').replace('_',' ').title()
                desc = table_desc.get(key, '')
                intro = f"<p class='expl'>{desc}</p>" if desc else ''
                table_htmls[key] = {"title": titre_tbl, "desc": table_desc.get(key,''), "html": html_table}
                continue
            else:
                # Table simple (toutes lignes) avec scroll si >20 et en-tête personnalisé
                # Remonter l'index en colonne pour contrôler l'en-tête
                if isinstance(tbl_fmt, pd.DataFrame):
                    tbl_fmt = tbl_fmt.reset_index()
                    # Ajout colonne Numéro sauf statut
                    if key != 'repartition_par_statut':
                        tbl_fmt.insert(0, 'Numéro', range(1, len(tbl_fmt) + 1))
                    # Renommer colonne catégorie (juste après Numéro si présente)
                    cat_col_idx = 1 if 'Numéro' in tbl_fmt.columns else 0
                    cat_col = tbl_fmt.columns[cat_col_idx]
                    tbl_fmt.rename(columns={cat_col: header_first_map.get(key, cat_col)}, inplace=True)
                # Construire manuellement la table HTML pour éviter double ligne d'en-tête
                headers = ''.join(f'<th>{h}</th>' for h in tbl_fmt.columns)
                body_rows = []
                for _, r in tbl_fmt.iterrows():
                    cells = ''.join(f'<td>{r[c]}</td>' for c in tbl_fmt.columns)
                    body_rows.append(f'<tr>{cells}</tr>')
                # Marquer spécialement le tableau des centres pour le masquer dans Word
                table_class = 'tbl centres-table-skip-word' if key == 'repartition_par_centre' else 'tbl'
                html_table_inner = f"<table class='{table_class}'><thead><tr>{headers}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"
                html_table = f"<div class='scroll-box'>{html_table_inner}</div>" if original_len>20 else html_table_inner
                if key == 'repartition_par_statut':
                    titre = 'Répartition par statut de traitement'
                else:
                    titre = key.replace('repartition_par_', 'Répartition par ').replace('_', ' ').title()
                desc = table_desc.get(key, '')
                intro = f"<p class='expl'>{desc}</p>" if desc else ''
                table_htmls[key] = {"title": titre, "desc": table_desc.get(key,''), "html": html_table}

    # Tableau d'évolution mensuelle (en plus du graphique)
    if isinstance(evo, pd.DataFrame) and not evo.empty:
        evo_tbl = evo.copy()
        # Normaliser index en colonne Mois
        if evo_tbl.index.name is None:
            evo_tbl = evo_tbl.reset_index().rename(columns={'index': 'Mois'})
        else:
            evo_tbl = evo_tbl.reset_index().rename(columns={evo_tbl.index.name: 'Mois'})
        # Reformater 'Mois' en libellés français complets (ex: janvier 2024)
        try:
            # evo a des clés type '2024-01'; parser
            def _fmt_month(m):
                try:
                    dt = pd.Period(m, freq='M').to_timestamp()
                except Exception:
                    return m
                mois_fr = ['janvier','février','mars','avril','mai','juin','juillet','août','septembre','octobre','novembre','décembre'][dt.month-1]
                return f"{mois_fr} {dt.year}"
            evo_tbl['Mois'] = evo_tbl['Mois'].apply(_fmt_month)
        except Exception:
            pass
        # Formater colonnes numériques
        for col in evo_tbl.columns:
            if col.lower().startswith('montant') or col in {'montant_total','sum'}:
                evo_tbl[col] = evo_tbl[col].apply(lambda x: fmt(x))
            elif col in {'nb_prestations','count'}:
                evo_tbl[col] = evo_tbl[col].apply(lambda x: fmt(x))
        # Ajouter Numéro
        evo_tbl.insert(0, 'Numéro', range(1, len(evo_tbl) + 1))
        # Renommer colonnes pour affichage pro
        evo_tbl.rename(columns={'montant_total': 'Montant total', 'nb_prestations': 'Nombre de prestations'}, inplace=True)
        headers = ''.join(f"<th>{h}</th>" for h in evo_tbl.columns)
        body_rows = []
        for _, r in evo_tbl.iterrows():
            cells = ''.join(f"<td>{r[c]}</td>" for c in evo_tbl.columns)
            body_rows.append(f"<tr>{cells}</tr>")
        html_table_inner = f"<table class='tbl'><thead><tr>{headers}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"
        html_table = f"<div class='scroll-box'>{html_table_inner}</div>" if len(evo_tbl) > 20 else html_table_inner
    desc = table_desc.get('evolution_mensuelle_tableau','')
    intro = f"<p class='expl'>{desc}</p>" if desc else ''
    table_htmls['evolution_mensuelle_tableau'] = {"title": 'Évolution mensuelle (tableau)', "desc": table_desc.get('evolution_mensuelle_tableau',''), "html": html_table}

    # Tableau dédié : Nombre total de prestations par mois (pour affichage clair)
    try:
        nb_tbl = evo.copy()
        # nb_tbl a index 'mois' et colonnes montant_total, nb_prestations
        nb_display = nb_tbl[['nb_prestations']].reset_index()
        # formatter mois lisible si nécessaire
        if nb_display.columns[0] != 'Mois':
            nb_display = nb_display.rename(columns={nb_display.columns[0]: 'Mois'})
        # ajouter numéro
        nb_display.insert(0, 'Numéro', range(1, len(nb_display) + 1))
        headers = ''.join(f"<th>{h}</th>" for h in nb_display.columns)
        body_rows = []
        for _, r in nb_display.iterrows():
            cells = ''.join(f"<td>{r[c]}</td>" for c in nb_display.columns)
            body_rows.append(f"<tr>{cells}</tr>")
        html_table_inner = f"<table class='tbl'><thead><tr>{headers}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"
        html_table_nb = f"<div class='scroll-box'>{html_table_inner}</div>" if len(nb_display) > 20 else html_table_inner
        table_htmls['nombre_prestations_par_mois'] = {"title": 'Nombre total de prestations par mois', "desc": 'Tableau synthétique du nombre de prestations observées mois par mois.', "html": html_table_nb}
    except Exception:
        pass

    # Analyse demandée : Nombre mensuel de chaque prestation (acte)
    try:
        acte_col = next((c for c in reps.get('repartition_par_acte', pd.DataFrame()).index.names if True), None)
    except Exception:
        acte_col = None
    # Better: detect column name for 'acte' in original dataframe via header_first_map keys
    try:
        # si la répartition par acte existe, utiliser son index
        if 'repartition_par_acte' in reps:
            actes_index = reps['repartition_par_acte'].index
            # construire un dataframe long mois x acte counts à partir de evo source (on a besoin du df original) -> utiliser df si accessible
            # fallback: si evo contient les mois seulement, reconstruire depuis le dataframe initial n'est pas trivial ici
            # donc tenter de recharger source et recalculer
            df_src = charger(SOURCE_FILE)
            date_col = next((c for c in df_src.columns if c.lower() == 'date'), None)
            acte_col_name = next((c for c in df_src.columns if c.lower() == 'acte'), None)
            if date_col and acte_col_name:
                tmp = df_src[[date_col, acte_col_name]].copy()
                tmp['Mois'] = tmp[date_col].dt.to_period('M').astype(str)
                tmp['Annee'] = tmp[date_col].dt.year
                
                # Séparer les données par année
                for annee in [2024, 2025]:
                    tmp_annee = tmp[tmp['Annee'] == annee].copy()
                    if tmp_annee.empty:
                        continue
                        
                    # Créer tableau pivot : actes en lignes, mois en colonnes
                    pivot_data = tmp_annee.groupby([acte_col_name, 'Mois']).size().unstack(fill_value=0)
                    
                    if pivot_data.empty:
                        continue
                    
                    # Formatter les noms de mois en français
                    def _fmt_m(m):
                        try:
                            dt = pd.Period(m, freq='M').to_timestamp()
                            mois_fr = ['janvier','février','mars','avril','mai','juin','juillet','août','septembre','octobre','novembre','décembre'][dt.month-1]
                            return f"{mois_fr}"
                        except Exception:
                            return m
                    
                    # Renommer les colonnes avec les mois français
                    pivot_data.columns = [_fmt_m(col) for col in pivot_data.columns]
                    
                    # Trier les actes par total décroissant
                    pivot_data['Total'] = pivot_data.sum(axis=1)
                    pivot_data = pivot_data.sort_values('Total', ascending=False)
                    
                    # Construire le tableau HTML
                    headers = '<th>Acte</th>' + ''.join(f'<th>{col}</th>' for col in pivot_data.columns)
                    body_rows = []
                    for acte, row in pivot_data.iterrows():
                        cells = [f'<td style="font-weight:600;">{acte}</td>']
                        for col in pivot_data.columns:
                            val = int(row[col]) if row[col] > 0 else '-'
                            if col == 'Total':
                                cells.append(f'<td style="font-weight:600;background:#f0f8ff;">{val}</td>')
                            else:
                                cells.append(f'<td>{val}</td>')
                        body_rows.append(f"<tr>{''.join(cells)}</tr>")
                    
                    html_table_inner = f"<table class='tbl'><thead><tr>{headers}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"
                    
                    # Créer des entrées séparées pour chaque année
                    table_key = f'nombre_mensuel_par_acte_{annee}'
                    table_htmls[table_key] = {
                        "title": f'Nombre mensuel de chaque prestation (acte) - {annee}', 
                        "desc": f'Tableau croisé dynamique pour l\'année {annee} : actes en lignes, mois en colonnes, avec totaux par acte.', 
                        "html": html_table_inner
                    }
                    
                    # Générer le graphique en lignes multiples correspondant
                    chart_filename = f'evolution_mensuelle_actes_{annee}.png'
                    chart_title = f'Évolution mensuelle des prestations par acte - {annee}'
                    chart_path = plot_monthly_lines(pivot_data, chart_title, chart_filename, annee)
                    if chart_path:
                        # Ajouter à la liste des images générées
                        images_paths.append(chart_path)
    except Exception:
        pass

    # Nouvelle analyse demandée : Montant moyen par prestation pour chaque acte (2024 vs 2025)
    try:
        print("[DEBUG] Début de l'analyse montant moyen par acte...")
        df_src = charger(SOURCE_FILE)
        print(f"[DEBUG] Données chargées: {len(df_src)} lignes")
        
        date_col = next((c for c in df_src.columns if c.lower() == 'date'), None)
        montant_col = next((c for c in df_src.columns if 'montant' in c.lower()), None)
        acte_col = next((c for c in df_src.columns if c.lower() == 'acte'), None)
        
        print(f"[DEBUG] Colonnes trouvées: date={date_col}, montant={montant_col}, acte={acte_col}")
        
        if date_col and montant_col and acte_col:
            print("[DEBUG] Toutes les colonnes sont disponibles, traitement...")
            tmp = df_src[[date_col, montant_col, acte_col]].copy()
            tmp['Annee'] = tmp[date_col].dt.year
            
            print(f"[DEBUG] Années disponibles: {sorted(tmp['Annee'].unique())}")
            
            # Calculer montant moyen par prestation pour chaque acte et chaque année
            stats_par_acte = []
            
            # Grouper par acte et année
            grouped = tmp.groupby([acte_col, 'Annee']).agg({
                montant_col: ['count', 'sum', 'mean']
            }).reset_index()
            
            # Aplatir les colonnes multi-niveaux
            grouped.columns = [acte_col, 'Annee', 'nb_prestations', 'montant_total', 'montant_moyen']
            
            print(f"[DEBUG] Données groupées: {len(grouped)} lignes")
            print(f"[DEBUG] Aperçu des données groupées:")
            print(grouped.head())
            
            # Restructurer pour avoir un tableau avec actes en lignes et années en colonnes
            actes_uniques = sorted(grouped[acte_col].unique())
            print(f"[DEBUG] Actes uniques trouvés: {actes_uniques}")
            
            tableau_data = []
            for acte in actes_uniques:
                row_data = {'Acte': acte}
                
                # Données pour 2024
                data_2024 = grouped[(grouped[acte_col] == acte) & (grouped['Annee'] == 2024)]
                if not data_2024.empty:
                    row_data['Prestations_2024'] = int(data_2024.iloc[0]['nb_prestations'])
                    row_data['Montant_moyen_2024'] = data_2024.iloc[0]['montant_moyen']
                else:
                    row_data['Prestations_2024'] = 0
                    row_data['Montant_moyen_2024'] = 0
                
                # Données pour 2025
                data_2025 = grouped[(grouped[acte_col] == acte) & (grouped['Annee'] == 2025)]
                if not data_2025.empty:
                    row_data['Prestations_2025'] = int(data_2025.iloc[0]['nb_prestations'])
                    row_data['Montant_moyen_2025'] = data_2025.iloc[0]['montant_moyen']
                else:
                    row_data['Prestations_2025'] = 0
                    row_data['Montant_moyen_2025'] = 0
                
                # Calculer l'évolution
                if row_data['Montant_moyen_2024'] > 0 and row_data['Montant_moyen_2025'] > 0:
                    evolution_pct = ((row_data['Montant_moyen_2025'] - row_data['Montant_moyen_2024']) / row_data['Montant_moyen_2024']) * 100
                    row_data['Evolution'] = evolution_pct
                else:
                    row_data['Evolution'] = None
                
                tableau_data.append(row_data)
            
            # Trier par montant moyen 2024 décroissant
            tableau_data.sort(key=lambda x: x['Montant_moyen_2024'], reverse=True)
            
            # Créer le tableau HTML
            if tableau_data:
                print(f"[DEBUG] Création du tableau HTML avec {len(tableau_data)} actes...")
                headers = [
                    'Acte', 
                    'Prestations 2024', 'Montant moyen 2024',
                    'Prestations 2025', 'Montant moyen 2025',
                    'Évolution (%)'
                ]
                
                # Construire les lignes du tableau
                body_rows = []
                for row in tableau_data:
                    # Formatage de l'évolution
                    if row['Evolution'] is not None:
                        evolution_val = row['Evolution']
                        if evolution_val > 0:
                            evolution_txt = f"↗️ +{evolution_val:.1f}%"
                            evolution_color = "green"
                        else:
                            evolution_txt = f"↘️ {evolution_val:.1f}%"
                            evolution_color = "red"
                    else:
                        evolution_txt = "-"
                        evolution_color = "gray"
                    
                    cells = [
                        f"<td style='font-weight:600;'>{row['Acte']}</td>",
                        f"<td>{fmt(row['Prestations_2024'])}</td>",
                        f"<td style='font-weight:600;'>{fmt(row['Montant_moyen_2024'])}</td>",
                        f"<td>{fmt(row['Prestations_2025'])}</td>",
                        f"<td style='font-weight:600;'>{fmt(row['Montant_moyen_2025'])}</td>",
                        f"<td style='font-weight:600;color:{evolution_color};'>{evolution_txt}</td>"
                    ]
                    body_rows.append(f"<tr>{''.join(cells)}</tr>")
                
                headers_html = ''.join(f'<th>{h}</th>' for h in headers)
                html_table_inner = f"<table class='tbl'><thead><tr>{headers_html}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"
                
                table_htmls['montant_moyen_par_acte'] = {
                    "title": 'Montant moyen par type de prestation : comparaison 2024 vs 2025', 
                    "desc": 'Analyse comparative du coût moyen de chaque type de prestation entre 2024 (12 mois) et 2025 (période partielle), avec calcul de l\'évolution par acte.', 
                    "html": html_table_inner
                }
                
                print("[DEBUG] Tableau HTML 'montant_moyen_par_acte' créé avec succès!")
                
                # Générer le graphique correspondant
                chart_filename = 'montant_moyen_par_acte_comparison.png'
                chart_title = 'Comparaison du montant moyen par type de prestation : 2024 vs 2025'
                print(f"Génération du graphique {chart_filename}...")
                chart_path = plot_montant_moyen_comparison(tableau_data, chart_title, chart_filename)
                if chart_path:
                    images_paths.append(chart_path)
                    print(f"Graphique généré : {chart_path}")
                else:
                    print("Erreur lors de la génération du graphique")
            else:
                print("[DEBUG] Aucune donnée pour le tableau montant_moyen_par_acte")
        else:
            print(f"[DEBUG] Colonnes manquantes pour l'analyse montant moyen par acte")
    except Exception as e:
        print(f"Erreur lors de l'analyse du montant moyen par acte: {e}")

    # Nouvelle analyse demandée : Proportion de chaque prestation (excepté pharmacie) sur le total
    try:
        df_src = charger(SOURCE_FILE)
        acte_col = next((c for c in df_src.columns if c.lower() == 'acte'), None)
        
        if acte_col:
            # Filtrer pour exclure la pharmacie
            df_no_pharma = df_src[~df_src[acte_col].str.lower().str.contains('pharmaci', na=False)].copy()
            
            # Compter les prestations par acte (sans pharmacie)
            counts_par_acte = df_no_pharma[acte_col].value_counts()
            total_prestations_no_pharma = len(df_no_pharma)
            
            # Calculer les proportions
            proportions_data = []
            for acte, count in counts_par_acte.items():
                proportion_pct = (count / total_prestations_no_pharma) * 100
                proportions_data.append({
                    'Acte': acte,
                    'Nombre_prestations': count,
                    'Proportion_pct': proportion_pct
                })
            
            # Trier par proportion décroissante
            proportions_data.sort(key=lambda x: x['Proportion_pct'], reverse=True)
            
            # Créer le tableau HTML
            if proportions_data:
                headers = [
                    'Acte', 
                    'Nombre de prestations',
                    'Proportion (%)',
                    'Représentation visuelle'
                ]
                
                # Construire les lignes du tableau
                body_rows = []
                for row in proportions_data:
                    # Barre de progression visuelle
                    bar_width = min(100, row['Proportion_pct'] * 2)  # Ajuster l'échelle pour la visibilité
                    progress_bar = f"<div style='background:#e0e0e0;width:100px;height:15px;border-radius:3px;'><div style='background:#4CAF50;width:{bar_width}px;height:15px;border-radius:3px;'></div></div>"
                    
                    cells = [
                        f"<td style='font-weight:600;'>{row['Acte']}</td>",
                        f"<td>{fmt(row['Nombre_prestations'])}</td>",
                        f"<td style='font-weight:600;'>{row['Proportion_pct']:.2f}%</td>",
                        f"<td>{progress_bar}</td>"
                    ]
                    body_rows.append(f"<tr>{''.join(cells)}</tr>")
                
                # Ajouter une ligne de total
                total_row = f"<tr style='background:#f8f9fa;border-top:2px solid #ddd;'><td style='font-weight:600;'>TOTAL (sans pharmacie)</td><td style='font-weight:600;'>{fmt(total_prestations_no_pharma)}</td><td style='font-weight:600;'>100.00%</td><td>-</td></tr>"
                body_rows.append(total_row)
                
                headers_html = ''.join(f'<th>{h}</th>' for h in headers)
                html_table_inner = f"<table class='tbl'><thead><tr>{headers_html}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"
                
                table_htmls['proportion_prestations_sans_pharmacie'] = {
                    "title": 'Proportion de chaque prestation (hors pharmacie) sur le total', 
                    "desc": f'Répartition proportionnelle des {total_prestations_no_pharma} prestations par type d\'acte, pharmacie exclue. La formule utilisée : (Nombre de prestations par acte / Total prestations sans pharmacie) × 100.', 
                    "html": html_table_inner
                }
                
                # Générer le graphique correspondant
                chart_filename = 'proportions_prestations_sans_pharmacie.png'
                chart_title = 'Proportions des prestations par type (hors pharmacie)'
                print(f"Génération du graphique {chart_filename}...")
                chart_path = plot_proportions_prestations(proportions_data, chart_title, chart_filename)
                if chart_path:
                    images_paths.append(chart_path)
                    print(f"Graphique généré : {chart_path}")
                else:
                    print("Erreur lors de la génération du graphique des proportions")
    except Exception:
        pass

    # ==== NOUVELLES ANALYSES DEMANDÉES ====
    try:
        # 1. Taux de recours aux prestations globales
        # 1. Taux de recours aux prestations globales
        # Calcul basé sur les données réelles des bénéficiaires
        
        # Calculer le nombre de mutualistes ayant consommé à partir du tableau des bénéficiaires
        mutualistes_ayant_consomme_reel = 0
        if 'repartition_par_beneficiaire' in reps and isinstance(reps['repartition_par_beneficiaire'], pd.DataFrame):
            beneficiaires_df = reps['repartition_par_beneficiaire']
            if 'mutualistes_distincts' in beneficiaires_df.columns:
                # Somme des adhérents + ayants droits (exclure les lignes calculées)
                for idx, row in beneficiaires_df.iterrows():
                    if isinstance(idx, str) and (idx.lower().startswith('adherent') or idx.lower().startswith('adhérent') or (idx.lower().startswith('ayant') and 'droit' in idx.lower())):
                        mutualistes_ayant_consomme_reel += int(row['mutualistes_distincts'])
        
        # Si on n'arrive pas à calculer depuis les bénéficiaires, on utilise la valeur globale
        if mutualistes_ayant_consomme_reel == 0:
            mutualistes_ayant_consomme_reel = global_indic.get('mutualistes_ayant_consomme', 0)
        
        # Configuration: Nombre total de mutualistes éligibles (valeur réelle fournie)
        TOTAL_MUTUALISTES_ELIGIBLES = 5284
        
        if mutualistes_ayant_consomme_reel > 0 and TOTAL_MUTUALISTES_ELIGIBLES > 0:
            taux_recours = (mutualistes_ayant_consomme_reel / TOTAL_MUTUALISTES_ELIGIBLES) * 100
            
            # Tableau pour le taux de recours
            recours_data = [
                ('Mutualistes ayant consommé (Adhérents + Ayants droits)', mutualistes_ayant_consomme_reel),
                ('Total mutualistes éligibles', TOTAL_MUTUALISTES_ELIGIBLES),
                ('Taux de recours (%)', round(taux_recours, 2))
            ]
            
            headers_recours = ['Indicateur', 'Valeur']
            body_rows_recours = []
            for label, value in recours_data:
                formatted_value = fmt(value) if isinstance(value, (int, float)) and not isinstance(value, bool) else str(value)
                body_rows_recours.append(f"<tr><td>{label}</td><td>{formatted_value}</td></tr>")
            
            headers_html_recours = ''.join(f"<th>{h}</th>" for h in headers_recours)
            html_table_recours = f"<table class='tbl'><thead><tr>{headers_html_recours}</tr></thead><tbody>{''.join(body_rows_recours)}</tbody></table>"
            
            table_htmls['taux_recours_global'] = {
                "title": 'Taux de recours aux prestations globales', 
                "desc": f'Indicateur mesurant la proportion de mutualistes ayant effectivement consommé des prestations sur l\'ensemble des mutualistes éligibles. Basé sur la somme des adhérents (1 229) et ayants droits (102) ayant consommé. Formule: (Nombre de mutualistes ayant consommé / Nombre total de mutualistes éligibles) × 100.', 
                "html": html_table_recours
            }
        
        # 2. Coût moyen de consommation par bénéficiaire 
        cout_moyen_beneficiaire = global_indic.get('cout_moyen_par_beneficiaire', 0)
        montant_total_consommation = global_indic.get('montant_total', 0)
        
        if cout_moyen_beneficiaire > 0:
            beneficiaire_data = [
                ('Montant total de consommation', fmt(montant_total_consommation)),
                ('Nombre de mutualistes ayant consommé', mutualistes_ayant_consomme_reel),
                ('Coût moyen par bénéficiaire', fmt(cout_moyen_beneficiaire))
            ]
            
            headers_beneficiaire = ['Indicateur', 'Valeur']
            body_rows_beneficiaire = []
            for label, value in beneficiaire_data:
                body_rows_beneficiaire.append(f"<tr><td>{label}</td><td>{value}</td></tr>")
            
            headers_html_beneficiaire = ''.join(f"<th>{h}</th>" for h in headers_beneficiaire)
            html_table_beneficiaire = f"<table class='tbl'><thead><tr>{headers_html_beneficiaire}</tr></thead><tbody>{''.join(body_rows_beneficiaire)}</tbody></table>"
            
            table_htmls['cout_moyen_consommation'] = {
                "title": 'Coût moyen de consommation par bénéficiaire', 
                "desc": f'Analyse du coût moyen supporté par chaque mutualiste ayant effectivement consommé des prestations. Basé sur {mutualistes_ayant_consomme_reel} mutualistes (adhérents + ayants droits). Formule: Montant total de consommation / Nombre de mutualistes ayant consommé.', 
                "html": html_table_beneficiaire
            }
            
    except Exception as e:
        print(f"Erreur lors du calcul des nouvelles analyses: {e}")

    # Tableau d'évolution trimestrielle
    if isinstance(evo, pd.DataFrame) and not evo.empty:
        # Reconstituer trimestre à partir de la période mensuelle
        try:
            evo_q = evo.copy()
            evo_q = evo_q.reset_index().rename(columns={evo_q.index.name or 'index': 'Mois'})
            # Convertir en période mensuelle puis en trimestre
            def _to_q(m):
                try:
                    p = pd.Period(m, freq='M')
                    q = p.asfreq('Q')
                    # Format Q1 2024 -> T1 2024
                    return f"T{q.quarter} {q.year}"
                except Exception:
                    return m
            evo_q['Trimestre'] = evo_q['Mois'].apply(_to_q)
            q_agg = evo_q.groupby('Trimestre')[['montant_total','nb_prestations']].sum().reset_index()
            # Trier par année/trimestre chronologiquement
            def _sort_key(t):
                try:
                    parts = t.split()
                    q = int(parts[0][1:])
                    y = int(parts[1])
                    return (y, q)
                except Exception:
                    return (9999, 99)
            q_agg = q_agg.sort_values(key=lambda s: s.apply(_sort_key))
            # Formater
            q_agg.insert(0, 'Numéro', range(1, len(q_agg)+1))
            for col in ['montant_total','nb_prestations']:
                if col in q_agg.columns:
                    q_agg[col] = q_agg[col].apply(lambda x: fmt(x))
            q_agg.rename(columns={'montant_total':'Montant total','nb_prestations':'Nombre de prestations'}, inplace=True)
            headers_q = ''.join(f"<th>{h}</th>" for h in q_agg.columns)
            body_rows_q = []
            for _, r in q_agg.iterrows():
                cells = ''.join(f"<td>{r[c]}</td>" for c in q_agg.columns)
                body_rows_q.append(f"<tr>{cells}</tr>")
            html_table_inner_q = f"<table class='tbl'><thead><tr>{headers_q}</tr></thead><tbody>{''.join(body_rows_q)}</tbody></table>"
            html_table_q = f"<div class='scroll-box'>{html_table_inner_q}</div>" if len(q_agg) > 20 else html_table_inner_q
            desc_q = table_desc.get('evolution_trimestrielle_tableau','')
            intro_q = f"<p class='expl'>{desc_q}</p>" if desc_q else ''
            table_htmls['evolution_trimestrielle_tableau'] = {"title": 'Évolution trimestrielle (tableau)', "desc": table_desc.get('evolution_trimestrielle_tableau',''), "html": html_table_q}
        except Exception:
            pass
    # Construction des blocs graphiques avec description inline
    graph_desc = {
        'montant_par_acte': (
            "Montant total par acte",
            "La hiérarchisation des actes par montant total met immédiatement en évidence les leviers prioritaires. Un pic isolé peut indiquer soit une catégorie réellement consommatrice, soit une dérive (tarif, codage). Le graphique sert donc autant d'instrument de pilotage que de déclencheur d'investigation qualitative. Il prépare la transition vers l'analyse de répartition relative."
        ),
        'montant_par_acte_pie': (
            "Répartition des montants par acte",
            "La structure relative des dépenses confirme ou tempère les impressions issues des montants absolus. Une forte concentration sur un petit nombre d'actes renforce la pertinence d'actions focalisées, tandis qu'une dispersion plus homogène appelle des mesures transversales (procédures, prévention globale). Le contraste barre / camembert enrichit le jugement managérial."
        ),
        'top10_centres': (
            "Top 10 centres – montants",
            "La concentration financière par centre signale les interlocuteurs clés pour la maîtrise du risque budgétaire. Un centre dominant peut refléter une spécialisation ou une dépendance commerciale. L'analyse aide à calibrer la fréquence des audits, la profondeur des revues médicales et la stratégie de contractualisation."
        ),
        'top10_centres_nbp': (
            "Top 10 centres – nombre de prestations",
            "Le classement en volume distingue les pôles d'activité intense. Comparé au classement financier, il révèle des divergences structurelles (fort volume mais faible panier moyen versus faible volume mais valeur élevée). Cette dissociation oriente différemment les actions : optimisation de flux d'un côté, contrôle des coûts unitaires de l'autre."
        ),
        'top10_partenaires': (
            "Top 10 partenaires – montants",
            "Les partenaires externes majeurs forment un périmètre de dépendance critique. Une concentration excessive expose à un risque opérationnel (capacité, rupture) ou tarifaire. La visualisation incite à diversifier ou à sécuriser les accords par clauses adaptées."
        ),
        'top10_partenaires_nbp': (
            "Top 10 partenaires – nombre de prestations",
            "Un partenaire sur‑sollicité en nombre mais modeste en montant est souvent le maillon logistique du parcours de soins courant. Sa performance influence directement l'expérience adhérent et la fluidité des remboursements. Préserver sa qualité devient stratégique."
        ),
        'top10_adherents_montant': (
            "Top 10 adhérents – montants",
            "La concentration sur quelques adhérents oriente vers des suivis individualisés : pathologies chroniques, épisodes lourds, potentiels cas de récurrence évitable. Cette vue nourrit la coordination médicale et les programmes de prévention ciblée à impact durable."
        ),
        'repartition_statut_pie': (
            "Répartition par statut de traitement",
            "La ventilation des statuts agit comme un indicateur de santé du processus administratif. Une dérive sur la proportion de dossiers rejetés ou en suspens peut entraîner un allongement des délais perçus et une érosion de confiance. Une baisse progressive des rejets témoigne de la maturité collaborative avec les prestataires."
        ),
        'evolution_mensuelle': (
            "Évolution mensuelle – graphique",
            "La combinaison courbe (montant) et barres (nombre) permet de repérer des divergences structurelles : montée des coûts unitaires (montant croît plus vite que le nombre) ou dilution (nombre progresse mais enveloppe stable). Ces signaux précoces permettent d'agir avant que l'effet cumulé ne pèse lourdement sur l'année."
        ),
        'evolution_trimestrielle': (
            "Évolution trimestrielle – graphique",
            "Le lissage trimestriel sépare tendance de bruit. Il aide à vérifier si une hausse récente est soutenable (tendance installée) ou transitoire (pic isolé). Cette lecture sert de pont avec les arbitrages budgétaires et les révisions de prévisions."
        ),
        'repartition_region_montant_nombre': (
            "Répartition géographique – régions",
            "La comparaison croisée des montants et du volume par région éclaire les modèles d'utilisation : forte intensité de recours vs forte valeur unitaire. Elle met en avant des profils contrastés exigeant des réponses distinctes (optimisation de l'accès, contrôle de la lourdeur des actes)."
        ),
        'repartition_province_montant_nombre': (
            "Répartition géographique – provinces",
            "Le détail provincial raffine les constats régionaux. Il permet d'isoler les foyers spécifiques de surconsommation ou de sous‑recours et de préparer des actions locales proportionnées (sensibilisation, renforcement de l'offre, audit ciblé)."
        ),
        'montant_par_type': (
            "Montant total par type",
            "Histogramme des montants cumulés par type : confirme la hiérarchie financière et sert de pont avec le détail des sous‑types."),
        'montant_par_sous_type': (
            "Montant total par sous-type (Top 30)",
            "Analyse concentrée sur les sous‑types les plus coûteux (Top 30 ou moins selon disponibilité) : met en lumière les niches spécifiques susceptibles d'actions ciblées (protocoles, prévention, renégociation)."),
        'pareto_actes': (
            "Courbe de Pareto des actes",
            "La courbe cumulative illustre quel pourcentage d'actes ou de catégories explique la majorité de la dépense. Un coude très précoce confirme la pertinence d'une stratégie focalisée sur un petit noyau; un profil plus étalé nécessite une approche plus systémique (processus, qualité globale)."
        ),
        'scatter_types': (
            "Dispersion actes (montant moyen vs fréquence)",
            "La matrice positionne chaque acte selon deux axes critiques : impact budgétaire unitaire et omniprésence. Le quadrant à forte fréquence et forte valeur suggère des cibles prioritaires où protocolisation, sensibilisation clinique ou renégociation peuvent générer un effet de levier maximal."
        ),
        'repartition_beneficiaire_pie': (
            "Répartition par bénéficiaire",
            "Cette vue distingue adhérent principal et ayants droit (ou autres statuts disponibles) en montrant leur part dans le nombre total de prestations. Elle met en lumière d'éventuelles concentrations d'utilisation et alimente des stratégies de prévention différenciées (éducation thérapeutique, suivi chronique, sensibilisation)."
        ),
        'repartition_genre_pie': (
            "Répartition par genre",
            "Le découpage par genre met en évidence des écarts potentiels de recours. Une sur‑représentation stable peut traduire des habitudes de consultation différentes ou des profils pathologiques distincts. Cet indicateur sert de base à des analyses d'équité et à l'adaptation de programmes de prévention ciblés." 
        ),
        'evolution_mensuelle_actes_2024': (
            "Évolution mensuelle des prestations par acte - 2024",
            "Le graphique en lignes multiples révèle les tendances temporelles et la saisonnalité de chaque type de prestation sur l'année 2024. Les pics simultanés sur plusieurs actes peuvent signaler des campagnes de santé, tandis que les évolutions divergentes orientent vers des analyses spécifiques par spécialité."
        ),
        'evolution_mensuelle_actes_2025': (
            "Évolution mensuelle des prestations par acte - 2025",
            "Cette visualisation pour 2025 permet la comparaison avec 2024 et l'identification des changements de comportement ou d'organisation. Les tendances émergentes servent à anticiper les besoins futurs et ajuster la stratégie de couverture en temps réel."
        ),
        'montant_moyen_par_acte_comparison': (
            "Comparaison du montant moyen par type de prestation : 2024 vs 2025",
            "Ce graphique en barres groupées visualise l'évolution des coûts moyens pour chaque type de prestation entre 2024 et 2025. Les différences de hauteur entre les barres révèlent les prestations qui ont connu une inflation ou une déflation tarifaire, guidant les décisions budgétaires et de couverture."
        ),
        'proportions_prestations_sans_pharmacie': (
            "Proportions des prestations par type (hors pharmacie)",
            "Ce graphique en barres horizontales présente visuellement la répartition proportionnelle des prestations par type d'acte, en excluant la pharmacie. Il permet d'identifier rapidement les prestations dominantes et d'analyser la structure des activités médicales pour une meilleure allocation des ressources."
        )
    }
    # Regrouper images par catégorie logique
    image_groups: dict[str, list[tuple[str,str,str]]] = {}
    for p in images_paths:
        stem = p.stem
        titre, desc = graph_desc.get(stem, (stem.replace('_',' ').title(), "Graphique de suivi."))
        img_tag = _img_tag_from_path(p)
        # Déduction catégorie
        if stem.startswith('montant_par_acte'):
            cat = 'repartition_par_acte'
        elif stem.startswith('top10_centres'):
            cat = 'repartition_par_centre'
        elif stem.startswith('top10_partenaires'):
            cat = 'repartition_par_partenaire'
        elif stem.startswith('repartition_statut'):
            cat = 'repartition_par_statut'
        elif stem.startswith('repartition_beneficiaire'):
            cat = 'repartition_par_beneficiaire'
        elif stem.startswith('repartition_genre'):
            cat = 'repartition_par_genre'
        elif stem.startswith('repartition_region'):
            cat = 'repartition_par_region'
        elif stem.startswith('repartition_province'):
            cat = 'repartition_par_province'
        elif stem.startswith('montant_par_type'):
            # Histogramme des montants par type : rattacher à la table repartition_par_type
            cat = 'repartition_par_type'
        elif stem.startswith('montant_par_sous_type'):
            # Histogramme des montants par sous-type : rattacher à la table repartition_par_sous_type
            cat = 'repartition_par_sous_type'
        elif stem.startswith('evolution_mensuelle_actes_2024'):
            # Nouveau graphique pour les prestations mensuelles par acte 2024
            cat = 'nombre_mensuel_par_acte_2024'
        elif stem.startswith('evolution_mensuelle_actes_2025'):
            # Nouveau graphique pour les prestations mensuelles par acte 2025
            cat = 'nombre_mensuel_par_acte_2025'
        elif stem.startswith('evolution_mensuelle'):
            cat = 'evolution_mensuelle'
        elif stem.startswith('evolution_trimestrielle'):
            cat = 'evolution_trimestrielle'
        elif stem.startswith('montant_moyen_par_acte_comparison'):
            # Graphique de comparaison des montants moyens par acte
            cat = 'montant_moyen_par_acte_comparison'
        elif stem.startswith('proportions_prestations_sans_pharmacie'):
            # Graphique des proportions de prestations (hors pharmacie)
            cat = 'proportion_prestations_sans_pharmacie'
        else:
            cat = stem  # fallback unique
        image_groups.setdefault(cat, []).append((titre, desc, img_tag))

    # Construction des blocs unifiés (table + graphiques)
    unified_blocks = []
    def combine_explanations(table_key: str, img_infos: list[tuple[str,str,str]]):
        tdesc = table_htmls.get(table_key, {}).get('desc','')
        gdesc = ' '.join([d for (_, d, _) in img_infos]) if img_infos else ''
        combined = (tdesc + ' ' + gdesc).strip()
        return combined
    # Liste des couples table->catégorie images à organiser
    couples = [
        ('repartition_par_acte','repartition_par_acte','Tableau : distribution des montants et fréquences par acte.','Graphiques : montants (barres) et part relative (camembert).'),
        ('repartition_par_centre','repartition_par_centre','Tableau : classement des centres (montants / prestations).','Graphiques : top montants et top fréquences.'),
        ('repartition_par_partenaire','repartition_par_partenaire','Tableau : classement des partenaires (montants / prestations).','Graphiques : top montants et top fréquences.'),
        ('repartition_par_statut','repartition_par_statut','Tableau : répartition quantitative par statut.','Graphique : ventilation proportionnelle.'),
        ('repartition_par_beneficiaire','repartition_par_beneficiaire','Tableau : répartition des prestations par catégorie de bénéficiaire.','Graphique : part relative de chaque catégorie.'),
        ('repartition_par_genre','repartition_par_genre','Tableau : répartition des prestations par genre.','Graphique : part relative par genre.'),
        ('repartition_par_type','repartition_par_type','Tableau : structure par grands types (montants & volumes).','Graphique : histogramme des montants par type.'),
        ('repartition_par_sous_type','repartition_par_sous_type','Tableau : détail des sous-types triés par montant.','Graphique : histogramme (Top 30 sous-types).'),
        ('repartition_par_region','repartition_par_region','Tableau : montants et volumes par région ordonnés par poids financier.','Graphique : comparaison visuelle montants (barres) et volumes (entre parenthèses).'),
        ('repartition_par_province','repartition_par_province','Tableau : détail provincial (ordre décroissant) avec regroupement régional.','Graphique : mise en avant des provinces les plus contributrices.'),
        ('evolution_mensuelle_tableau','evolution_mensuelle','Tableau : détail chronologique mois par mois (montants et volumes).','Graphique : dynamique conjointe montants (courbe) et volumes (barres).'),
        ('evolution_trimestrielle_tableau','evolution_trimestrielle','Tableau : consolidation trimestrielle (lissage des fluctuations mensuelles).','Graphique : trajectoire structurée des montants agrégés par trimestre.'),
        ('nombre_mensuel_par_acte_2024','nombre_mensuel_par_acte_2024','Tableau : évolution mensuelle des prestations par acte en 2024.','Graphique : courbes multiples montrant les tendances temporelles pour chaque type de prestation.'),
        ('nombre_mensuel_par_acte_2025','nombre_mensuel_par_acte_2025','Tableau : évolution mensuelle des prestations par acte en 2025.','Graphique : courbes multiples permettant la comparaison avec 2024 et l\'analyse des tendances émergentes.'),
        ('montant_moyen_par_acte','montant_moyen_par_acte_comparison','Tableau : montants moyens par type de prestation comparés entre 2024 et 2025.','Graphique : barres groupées visualisant l\'évolution des coûts moyens pour identifier les inflations tarifaires par spécialité.'),
        ('proportion_prestations_sans_pharmacie','proportion_prestations_sans_pharmacie','Tableau : répartition proportionnelle des prestations par acte (hors pharmacie).','Graphique : barres horizontales visualisant la dominance relative de chaque type de prestation dans l\'activité globale.'),
        ('taux_recours_global','taux_recours_global','Tableau : indicateurs de taux de recours aux prestations.','Analyse de la participation effective des mutualistes aux prestations disponibles.'),
        ('cout_moyen_consommation','cout_moyen_consommation','Tableau : coût moyen de consommation par bénéficiaire.','Évaluation de l\'effort financier moyen par mutualiste ayant consommé.')
    ]
    used_image_cats = set()
    for table_key, img_cat, intro_table, intro_graph in couples:
        print(f"[DEBUG] Vérification couple: table='{table_key}', image_cat='{img_cat}'")
        print(f"[DEBUG] Table existe: {table_key in table_htmls}")
        print(f"[DEBUG] Image category existe: {img_cat in image_groups}")
        
        if table_key in table_htmls and img_cat in image_groups:
            print(f"[DEBUG] Couple OK: {table_key} -> {img_cat}")
            title = table_htmls[table_key]['title']
            combined_expl = combine_explanations(table_key, image_groups[img_cat])
            block = f"<section class='bloc'><h3>{title}</h3><p class='expl'>{combined_expl}</p><p class='intro'>{intro_table}</p>{table_htmls[table_key]['html']}<p class='intro'>{intro_graph}</p>" + '\n'.join(
                f"<figure>{img_tag}<figcaption>{t}</figcaption></figure>" for t,d,img_tag in image_groups[img_cat]
            ) + "</section>"
            unified_blocks.append(block)
            used_image_cats.add(img_cat)

    # Ajouter tables restantes sans images groupées
    print(f"[DEBUG] Vérification tables restantes...")
    for k, meta in table_htmls.items():
        print(f"[DEBUG] Table '{k}': dans couples={any(k == c[0] for c in couples)}")
        if any(k == c[0] for c in couples):
            continue
        # Inclure explicitement nos tables mensuelles personnalisées (sauf celles déjà dans couples)
        if k.endswith('_tableau') or k in {'repartition_par_region','repartition_par_province','repartition_par_type','repartition_par_sous_type','nombre_prestations_par_mois','montant_moyen_par_acte','taux_recours_global','cout_moyen_consommation'}:
            print(f"[DEBUG] Ajout table restante: {k}")
            unified_blocks.append(f"<section class='bloc'><h3>{meta['title']}</h3><p class='expl'>{meta['desc']}</p>{meta['html']}</section>")

    # Graphiques orphelins (sans tableau associé)
    for cat, infos in image_groups.items():
        # Ne pas exclure systématiquement les catégories de couples : seulement celles déjà utilisées
        if cat in used_image_cats:
            continue
        if cat in {'evolution_mensuelle','evolution_trimestrielle','pareto_actes','scatter_types'}:
            combined_desc = ' '.join([d for _, d, _ in infos])
            title = infos[0][0]
            figs = ''.join(f"<figure>{img}<figcaption>{t}</figcaption></figure>" for t,d,img in infos)
            unified_blocks.append(f"<section class='bloc'><h3>{title}</h3><p class='expl'>{combined_desc}</p>{figs}</section>")

    unified_html = '\n'.join(unified_blocks)
    objectifs_html = """
        <section>
            <h2>Introduction générale</h2>
            <p>Ce rapport analyse en profondeur les prestations de santé prises en charge par la Mutuelle sur la période allant de janvier 2024 à juillet 2025. Il vise à transformer un volume hétérogène de données opérationnelles en lecture stratégique exploitable pour la gouvernance, la maîtrise des risques médico‑financiers et l’orientation des actions correctrices. La progression des charges de santé, la variabilité des pratiques de facturation, la pression budgétaire sur la soutenabilité des régimes solidaires et l’attente de transparence des adhérents imposent un pilotage analytique précis et régulier.</p>
            <h3>Contexte</h3>
            <p>La Mutuelle évolue dans un environnement marqué par une progression continue des charges de santé, une variabilité des pratiques de facturation et une pression croissante sur la soutenabilité des régimes solidaires. Dans ce cadre, une analyse approfondie des prestations enregistrées entre janvier 2024 et juillet 2025 s’impose afin de fournir à la gouvernance une lecture claire et fiable des données. Cette démarche répond également à une attente de transparence de la part des adhérents et vise à renforcer le pilotage analytique des risques médico‑financiers.</p>
            <h3>Objectifs analytiques</h3>
            <ol>
                <li>Établir une photographie fiable des volumes, montants globaux, dispersion et coûts unitaires.</li>
                <li>Hiérarchiser les contributeurs majeurs : actes, types, sous types, centres, partenaires, zones.</li>
                <li>Déceler les signaux d’évolution : accélérations, inflexions, points de rupture (mensuel & trimestriel).</li>
                <li>Identifier les poches de concentration (top adhérents, nœuds géographiques, sous types émergents).</li>
                <li>Apprécier l’équilibre bénéficiaires / genre pour une lecture d’équité d’accès.</li>
                <li>Évaluer la qualité du flux de traitement via la distribution des statuts.</li>
                <li>Préparer une base rationnelle pour prioriser audits, renégociations ou actions de prévention ciblées.</li>
            </ol>
            <h3>Périmètre et fiabilisation des données</h3>
            <ul>
                <li><strong>Inclusion :</strong> prestations remboursables enregistrées (janv. 2024 – juil. 2025).</li>
                <li><strong>Normalisations :</strong> typographie des centres, homogénéisation des statuts, formats monétaires.</li>
                <li><strong>Correctif calendaire :</strong> dates août–décembre 2025 rebasculées sur 2024 (anomalie corrigée et tracée).</li>
                <li><strong>Qualité :</strong> élimination des doublons (extraction conservatoire), contrôle des colonnes critiques, conversion explicite des montants.</li>
                <li><strong>Exclusions :</strong> pas de modélisation prédictive avancée, ni benchmarking externe, ni analyse médicale qualitative.</li>
            </ul>
            <h3>Méthodologie synthétique</h3>
            <ul>
                <li>Agrégations multi axes (acte, type, sous type, centre, partenaire, région, province).</li>
                <li>Mesures descriptives : montants totaux, volumes, moyennes, médianes, dispersion implicite.</li>
                <li>Répartition relative (camemberts, histogrammes triés) pour traduire la concentration.</li>
                <li>Lecture temporelle double : mensuelle (sensibilité fine) et trimestrielle (tendance).</li>
                <li>Segmentation bénéficiaires / genre pour détecter d’éventuels déséquilibres.</li>
                <li>Visualisations annotées pour réduire l’ambiguïté interprétative.</li>
            </ul>
            <h3>Définitions opérationnelles (extraits)</h3>
            <ul>
                <li><strong>Montant total :</strong> somme des montants enregistrés.</li>
                <li><strong>Nombre de prestations :</strong> enregistrements unitaires après nettoyage.</li>
                <li><strong>Coût moyen par prestation :</strong> montant total / nombre de prestations.</li>
                <li><strong>Statuts :</strong> catégories de traitement (accepté / autre) harmonisées.</li>
                <li><strong>Type / sous type :</strong> niveaux hiérarchiques internes de classification.</li>
            </ul>
            <h3>Limites et prudences</h3>
            <ul>
                <li>Données = historique comptabilisé (non‑recours non mesuré).</li>
                <li>Pas d’ajustement clinique (complexité / sévérité non intégrées).</li>
                <li>Montants non ajustés pour inflation intra période.</li>
                <li>Retards possibles de facturation pouvant lisser certaines hausses.</li>
            </ul>
        </section>
        """
    explications_html = ""  # remplacé par descriptions inline
    # Section Traitement des données
    montant_paye = global_indic.get('montant_paye_total', '')
    montant_non_paye = global_indic.get('montant_non_paye', '')
    pct_paye = global_indic.get('pourcentage_paye_pct', '')
    nb_dates_corr = 0
    if 'DATES_CORRIGEES_DF' in globals() and DATES_CORRIGEES_DF is not None:
        nb_dates_corr = len(DATES_CORRIGEES_DF)
    traitement_rows = []
    def _fmt_brut(v):
        if isinstance(v,(int,float)) and not isinstance(v,bool):
            if float(v).is_integer():
                return f"{int(v):,}".replace(',', ' ')
            return f"{v:,.2f}".replace(',', ' ').replace('.', ',')
        return v
    # (Montant payé / non payé retirés sur demande)
    if nb_dates_corr:
        traitement_rows.append(f"<tr><td>Dates corrigées (août–déc. 2025 =&gt; 2024)</td><td>{nb_dates_corr}</td></tr>")
    if DUPLICATES_INFO:
        traitement_rows.append(f"<tr><td>Lignes doublons supprimées</td><td>{DUPLICATES_INFO.get('nb_doublons_lignes',0)}</td></tr>")
    traitement_html = f"""
    <section>
      <h2>1. Traitement des données</h2>
      <p>Opérations appliquées avant analyse :</p>
      <ul>
        <li>Normalisation des noms de colonnes / centres (majuscules, espaces).</li>
        <li>Conversion des dates, filtrage sur 2024–2025.</li>
        <li>Correction : toutes les dates d'août à décembre 2025 ont leur année remplacée par 2024.</li>
        <li>Nettoyage des montants (espaces / séparateurs) et conversion numérique.</li>
        <li>Renommage de « validite » en « statut » + uniformisation (minuscules).</li>
        <li>Suppression des doublons (lignes identiques) après extraction vers feuilles dédiées.</li>
        <li>Calcul montants payés : somme des montants où statut ∈ {"accepté","accepte"}; non payé = total - payé.</li>
      </ul>
      <table class='tbl'><thead><tr><th>Élément</th><th>Valeur</th></tr></thead><tbody>{''.join(traitement_rows) if traitement_rows else '<tr><td>(aucun)</td><td>-</td></tr>'}</tbody></table>
    <!-- Ligne de formules supprimée sur demande -->
    </section>
    """
    # Logo base64 si dispo
    logo_html = ''
    if LOGO_PATH.exists():
        try:
            with open(LOGO_PATH, 'rb') as lf:
                _b64 = base64.b64encode(lf.read()).decode('utf-8')
            logo_html = f"<img src='data:image/png;base64,{_b64}' alt='Logo' style='height:70px;float:right;margin-left:15px;' />"
        except Exception:
            logo_html = ''
    # Logo pour le footer (version réduite)
    footer_logo_html = ''
    if LOGO_PATH.exists():
        try:
            with open(LOGO_PATH, 'rb') as lf2:
                _b64_footer = base64.b64encode(lf2.read()).decode('utf-8')
            footer_logo_html = f"<img src='data:image/png;base64,{_b64_footer}' alt='Logo' style='height:34px;display:inline-block;' />"
        except Exception:
            footer_logo_html = ''
    html = f"""<!DOCTYPE html><html lang='fr'><head><meta charset='utf-8'><title>Rapport Prestations</title>
<style>
body{{font-family:Arial,Helvetica,sans-serif;margin:25px;}}
h1{{color:#2E4053;}}
table.tbl{{border-collapse:collapse;margin:10px 0;font-size:13px;}}
table.tbl th,table.tbl td{{border:1px solid #ccc;padding:4px 8px;text-align:left;}}
thead th{{background:#f2f2f2;}}
.grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(320px,1fr));gap:20px;}}
figure{{margin:0;}}
figcaption{{font-size:12px;text-align:center;margin-top:4px;color:#555;}}
.scroll-box{{max-height:480px;overflow-y:auto;border:1px solid #bbb;padding:4px;background:#fafafa;margin-top:6px;}}
details summary{{cursor:pointer;font-weight:600;margin:6px 0;}}
</style></head><body>
<div style='overflow:auto;'> {logo_html}<h1 style='margin-top:10px;'>Rapport Analytique des Prestations</h1></div>
{objectifs_html}
{traitement_html}
<h2>3. Indicateurs globaux</h2>
<p class='expl'>Tableau de synthèse générale (périmètre, dispersion, volumes et période d'observation) servant de point d'entrée à l'analyse.</p>
<table class='tbl'><thead><tr><th>Indicateur</th><th>Valeur</th></tr></thead><tbody>{rows_indic}</tbody></table>
<h2>4. Analyses détaillées</h2>
{unified_html}
<hr style='margin:40px 0 15px;border:none;border-top:1px solid #ccc;'>
<section>
    <h2>Conclusion</h2>
    <p>Cette analyse a permis de transformer des données brutes en une lecture stratégique, mettant en lumière les principaux contributeurs aux dépenses, les déséquilibres potentiels et les tendances d’évolution. Les constats dégagés ouvrent la voie à des actions concrètes : renforcer la maîtrise médico-financière, cibler les audits, optimiser les partenariats et anticiper les risques. L’enjeu est désormais de traduire ces enseignements en décisions opérationnelles afin de soutenir la soutenabilité du régime et d’outiller la gouvernance dans son pilotage stratégique.</p>
</section>
<footer style='text-align:center;font-size:13px;color:#555;font-style:italic;'>
    <div style='display:flex;align-items:center;justify-content:center;gap:10px;'>
        {footer_logo_html}
        <span>Mutuelle de la Police Nationale</span>
    </div>
</footer>
</body></html>"""
    with open('rapport_prestations.html','w',encoding='utf-8') as f:
        f.write(html)
    global LAST_HTML_REPORT
    LAST_HTML_REPORT = html


if __name__ == '__main__':
    main()
